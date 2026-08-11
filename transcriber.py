"""Desgrabador audiovisual → nota web (Wix) + reel a Facebook e Instagram.

Pensado para correr en la NUBE (GitHub Actions), disparado por un Google Apps Script
cuando un colaborador sube un video (o una SUBCARPETA con video + fotos + texto) a la
carpeta de Drive «videos notas actualidad». NO usa tokens de Claude: la desgrabación la
hace Gemini (gratis), que recibe el VIDEO COMPLETO (audio + texto en pantalla + subtítulos
+ imágenes) más el contexto adjunto.

Dos etapas (flujo CON revisión):

  ETAPA 1 — `run_transcribe_video(file, uploader)`  (al subir):
    1. junta los adjuntos de la subcarpeta (fotos + texto) como contexto
    2. Gemini desgraba el video → {hay_noticia, volanta, titulo, texto, resumen, mejor_momento_seg}
    3. saca la foto de portada en el segundo más representativo que indica Gemini
    4. arma el reel vertical 9:16 (si no hay noticia, recortado a 1 min) y lo sube a un Release
    5. SI HAY NOTICIA: crea la nota como BORRADOR en Wix (foto + video nativo) y avisa
       SI NO HAY: NO crea nota web; deja listo solo el reel y avisa para decidir
    6. registra la fila de contabilidad

  ETAPA 2 — `run_publish_video(file)`  (al mover el video a APROBADAS):
    - Con noticia: publica la nota web + reel a FB/IG con el resumen de caption.
    - Sin noticia: la web queda SUSPENDIDA; sale SOLO el reel (sin texto).
"""
import json
import re
import smtplib
import ssl
import time
from datetime import datetime
from email.message import EmailMessage
from email.utils import formataddr
from html import escape as _hesc
from pathlib import Path
from urllib.parse import quote

import requests

from platforms import facebook, instagram, wix
from utils.config import get
from utils.gemini import transcribe_to_nota
from utils.logger import get_logger
from utils.video_host import upload_reel
from video import best_parts_clip, duration_seconds, frame_at, remux_mp4, to_vertical_reel

logger = get_logger("transcriber")

LEDGER = Path(__file__).parent / ".videos_contabilidad.json"
WORK_DIR = Path(__file__).parent / "videos_preview"
VIDEO_EXTS = {".mp4", ".mov", ".mkv", ".avi", ".webm", ".m4v", ".mpg", ".mpeg"}
IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".webp"}
TEXT_EXTS = {".txt", ".md", ".docx"}
REEL_MAX_SIN_NOTICIA = 60  # segundos: tope del reel cuando no se pudo desgrabar


# ── Helpers de entorno ────────────────────────────────────────────────────────
def _site() -> str:
    return get("STORY_SITE_URL") or "www.diariolacampaña.com.ar"


def _platforms() -> list[str]:
    raw = get("STORIES_PLATFORMS") or "instagram,facebook"
    return [p.strip().lower() for p in raw.split(",") if p.strip()]


def _videos_folder() -> Path:
    return Path(get("VIDEOS_FOLDER") or (Path(__file__).parent / "videos"))


def _find_video(name: str, base: Path | None = None) -> Path | None:
    """Ubica el video bajado de Drive por nombre (en la raíz o en una subcarpeta);
    si no, agarra el más nuevo. `base` = carpeta donde buscar (default: la del diario;
    el desgrabador de la radio pasa su propia carpeta)."""
    folder = base or _videos_folder()
    if not folder.exists():
        return None
    if name:
        cand = folder / name
        if cand.exists():
            return cand
        for p in folder.rglob("*"):
            if p.is_file() and p.name == name:
                return p
    vids = [p for p in folder.rglob("*") if p.is_file() and p.suffix.lower() in VIDEO_EXTS]
    return max(vids, key=lambda p: p.stat().st_mtime) if vids else None


def _slug(s: str) -> str:
    s = re.sub(r"[^a-zA-Z0-9]+", "-", (s or "")).strip("-").lower()
    return s[:40] or "reel"


IMG_EXTS = {".jpg", ".jpeg", ".png", ".webp", ".gif"}


def _find_folder(name: str, base: Path | None = None) -> Path | None:
    """Ubica la SUBCARPETA de una nota-placa (Word + foto, sin video) por nombre.
    `base` = carpeta raíz donde buscar (default: la del diario)."""
    base = base or _videos_folder()
    if not base.exists():
        return None
    if name:
        cand = base / name
        if cand.is_dir():
            return cand
        for p in base.rglob("*"):
            if p.is_dir() and p.name == name:
                return p
    return None


def _parse_word(path: Path) -> tuple[str, str, list[str]]:
    """(volanta, titular, cuerpo) desde un .docx o .txt. 1er párrafo corto = volanta."""
    if path.suffix.lower() == ".txt":
        paras = [ln.strip() for ln in path.read_text(encoding="utf-8-sig", errors="ignore").splitlines() if ln.strip()]
    else:
        from docx import Document
        paras = [p.text.strip() for p in Document(str(path)).paragraphs if p.text.strip()]
    if not paras:
        return "", "", []
    es_vol = len(paras) >= 2 and (len(paras[0]) <= 55 or len(paras[0].split()) <= 7)
    if es_vol:
        return paras[0], paras[1], paras[2:]
    return "", paras[0], paras[1:]


# ── Adjuntos de la subcarpeta (fotos + texto de contexto) ─────────────────────
def _leer_texto(path: Path) -> str:
    try:
        if path.suffix.lower() == ".docx":
            from docx import Document
            return "\n".join(p.text for p in Document(str(path)).paragraphs if p.text.strip())
        return path.read_text(encoding="utf-8-sig", errors="ignore")
    except Exception as e:
        logger.warning(f"No se pudo leer el contexto {path.name}: {e}")
        return ""


def _recolectar_adjuntos(video: Path) -> tuple[str, list[Path]]:
    """Si el video está en una SUBCARPETA (no en la raíz de videos/), junta las fotos y
    textos hermanos como contexto. En la raíz no junta nada (evita mezclar notas)."""
    folder = video.parent
    try:
        if folder.resolve() == _videos_folder().resolve():
            return "", []
    except Exception:
        return "", []
    textos, imgs = [], []
    for p in sorted(folder.iterdir()):
        if not p.is_file() or p == video:
            continue
        if p.name.lower() == "contexto.txt":
            continue  # lo parsea _leer_contexto() aparte (datos del corresponsal)
        ext = p.suffix.lower()
        if ext in TEXT_EXTS:
            t = _leer_texto(p)
            if t.strip():
                textos.append(t.strip())
        elif ext in IMAGE_EXTS:
            imgs.append(p)
    if textos or imgs:
        logger.info(f"Adjuntos en «{folder.name}»: {len(textos)} texto(s), {len(imgs)} foto(s)")
    return "\n\n".join(textos), imgs


# ── Contexto del corresponsal (contexto.txt que deja el bot de WhatsApp) ──────
FIRMA_DEFAULT = ("Material enviado por un colaborador de la Red de Corresponsales — "
                 "Diario La Campaña · Radio del Centro")
_CTX_KEYS = ("origen", "nombre", "celular", "lugar", "descripcion", "autorizacion")


def _firma_texto() -> str:
    return (get("CORRESPONSALES_FIRMA") or FIRMA_DEFAULT).strip()


def _leer_contexto(folder: Path) -> dict | None:
    """Si en la subcarpeta del video hay un `contexto.txt` con los datos del corresponsal
    (lo escribe el webhook de WhatsApp), lo parsea a un dict con las claves de `_CTX_KEYS`.
    Formato `CLAVE: valor` (DESCRIPCION puede ser multilínea). Devuelve None si no existe o
    no tiene el marcador ORIGEN."""
    try:
        cand = folder / "contexto.txt"
        if not cand.exists():
            return None
        raw = cand.read_text(encoding="utf-8-sig", errors="ignore")
    except Exception as e:
        logger.warning(f"No se pudo leer contexto.txt: {e}")
        return None
    datos: dict[str, str] = {}
    actual = None
    for linea in raw.splitlines():
        m = re.match(r"\s*([A-Za-zÁÉÍÓÚÑ]+)\s*:\s*(.*)$", linea)
        if m and m.group(1).strip().lower() in _CTX_KEYS:
            actual = m.group(1).strip().lower()
            datos[actual] = m.group(2).strip()
        elif actual:  # continuación de un valor multilínea (ej. DESCRIPCION)
            datos[actual] = (datos[actual] + "\n" + linea.rstrip()).strip()
    if not datos.get("origen"):
        return None
    return datos


# ── Ledger de contabilidad ────────────────────────────────────────────────────
def _leer_ledger(path: Path | None = None) -> list[dict]:
    ledger = path or LEDGER
    if not ledger.exists():
        return []
    try:
        return list(json.loads(ledger.read_text(encoding="utf-8-sig")))
    except Exception:
        return []


def _guardar_ledger(rows: list[dict], path: Path | None = None) -> None:
    (path or LEDGER).write_text(json.dumps(rows, ensure_ascii=False, indent=2), encoding="utf-8")


def _buscar_fila(rows: list[dict], file: str) -> dict | None:
    for row in rows:
        if row.get("file") == file:
            return row
    return None


# ── Aviso por mail ────────────────────────────────────────────────────────────
def _boton(url: str, texto: str, color: str = "#e2620c") -> str:
    return (f'<a href="{url}" style="display:inline-block;background:{color};color:#fff;'
            f'text-decoration:none;padding:12px 20px;border-radius:6px;font-family:Arial;'
            f'font-size:16px;margin:6px 6px 6px 0">{texto}</a>')


def _html_aviso(intro_html: str, name: str, reel_url: str, draft_id: str, hay: bool) -> str:
    """Arma el cuerpo HTML del aviso con los botones (si hay APPROVE_WEBAPP_URL)."""
    webapp = get("APPROVE_WEBAPP_URL")
    tok = get("WEBAPP_TOKEN")
    t = f"&token={quote(tok)}" if tok else ""
    botones = ""
    if webapp:
        botones += _boton(f"{webapp}?action=approve&name={quote(name)}{t}", "✅ Aprobar y publicar")
        if reel_url:
            # Reproduce el reel dentro del navegador (web app) en vez de descargarlo.
            botones += _boton(f"{webapp}?action=preview&url={quote(reel_url)}{t}",
                              "👁️ Previsualizar video", color="#444")
        if hay and draft_id:
            botones += _boton(f"{webapp}?action=edit&draft={draft_id}{t}", "✏️ Corregir texto", color="#444")
            botones += _boton(f"{webapp}?action=delete&post={quote(draft_id)}{t}", "🗑️ Borrar borrador", color="#b00020")
    elif reel_url:
        botones += _boton(reel_url, "👁️ Ver el reel", color="#444")
    return (f'<div style="font-family:Arial;max-width:600px;color:#222;font-size:16px">'
            f'{intro_html}<div style="margin:22px 0">{botones}</div>'
            f'<p style="color:#777;font-size:13px">Si no ves los botones, aprobá moviendo el '
            f'video a la subcarpeta APROBADAS en Drive.</p></div>')


def _botones_foto(name: str, draft_id: str, reel_url: str) -> str:
    """Botones de revisión para los mails de FOTO (aprobar + previsualizar reel + corregir texto +
    borrar), iguales que en el mail de video. Reusa las acciones del endpoint web (`aprobar-video`):
    approve/preview/edit/delete. Corregir y borrar necesitan un borrador de Wix (draft_id);
    previsualizar necesita la URL del reel."""
    webapp = get("APPROVE_WEBAPP_URL"); tok = get("WEBAPP_TOKEN")
    if not webapp:
        return _boton(reel_url, "👁️ Ver el reel", color="#444") if reel_url else ""
    t = f"&token={quote(tok)}" if tok else ""
    b = _boton(f"{webapp}?action=approve&name={quote(name)}&kind=folder{t}", "✅ Aprobar y publicar")
    if reel_url:
        b += _boton(f"{webapp}?action=preview&url={quote(reel_url)}{t}", "👁️ Previsualizar reel", color="#444")
    if draft_id:
        b += _boton(f"{webapp}?action=edit&draft={quote(draft_id)}{t}", "✏️ Corregir texto", color="#444")
        b += _boton(f"{webapp}?action=delete&post={quote(draft_id)}{t}", "🗑️ Borrar borrador", color="#b00020")
    return b


def _reel_preview(fotos, slug: str) -> str:
    """Arma el reel de la/s foto/s y lo sube para poder PREVISUALIZARLO en la revisión (best-effort).
    Devuelve la URL o "" si falla (el mail sale sin ese botón). El reel definitivo se rearma igual al
    publicar; como la foto no lleva texto quemado, corregir el texto después no cambia el reel."""
    try:
        from video import foto_a_reel
        WORK_DIR.mkdir(exist_ok=True)
        reel_local = foto_a_reel(fotos, WORK_DIR / f"prev_{slug}.mp4", overlay=False)
        return upload_reel(reel_local)
    except Exception as e:  # noqa: BLE001
        logger.warning(f"No pude armar el reel de previsualización ({e}); el mail va sin ese botón.")
        return ""


def _enviar_aviso(asunto: str, cuerpo: str, html: str | None = None, destino: str = "") -> None:
    """Manda un mail al diario (reusa el SMTP del mailer). Best-effort. Si se pasa `html`,
    va como alternativa HTML (con botones). `destino` sobreescribe el destinatario (lo usa
    el desgrabador de la radio con VIDEOS_RADIO_NOTIFY_EMAIL)."""
    remitente = get("MAIL_FROM")
    password = get("MAIL_APP_PASSWORD")
    destino = (destino or "").strip() or get("VIDEOS_NOTIFY_EMAIL") or remitente
    if not remitente or not password or not destino:
        logger.warning("Sin credenciales de mail (MAIL_FROM/MAIL_APP_PASSWORD): no se manda el aviso.")
        return
    host = get("SMTP_HOST") or "smtp.gmail.com"
    port = int(get("SMTP_PORT") or 587)
    nombre_from = get("MAIL_FROM_NAME") or "Diario La Campaña"
    msg = EmailMessage()
    msg["From"] = formataddr((nombre_from, remitente))
    msg["To"] = destino
    msg["Subject"] = asunto
    msg.set_content(cuerpo)
    if html:
        msg.add_alternative(html, subtype="html")
    try:
        ctx = ssl.create_default_context()
        with smtplib.SMTP(host, port, timeout=60) as server:
            server.starttls(context=ctx)
            server.login(remitente, password)
            server.send_message(msg)
        logger.info(f"Aviso enviado a {destino}")
    except Exception as e:
        logger.error(f"No se pudo enviar el aviso por mail: {e}")


def _descargar(url: str, destino: Path) -> Path:
    r = requests.get(url, headers={"User-Agent": "Mozilla/5.0"}, timeout=180)
    r.raise_for_status()
    destino.write_bytes(r.content)
    return destino


def _avisar_nombre_repetido(nombre: str, fila: dict, kind: str = "nota") -> None:
    """Cuando una etapa-1 (video o foto-nota) se saltea porque el NOMBRE ya existe en el ledger
    (ya procesado), avisa por mail. Antes se quedaba MUDO y el colaborador no se enteraba de que
    su nota nueva se descartó por REUSAR un nombre ya usado (pasó 2026-07-27 con la carpeta «Moto»).

    Salvaguarda anti-aviso-falso: solo avisa si la entrada previa NO es de los últimos ~15 min.
    Así un doble disparo del MISMO ítem (mismo folder/archivo procesado recién en esta tanda) no
    genera un aviso falso; un nombre reusado de horas/días atrás sí lo genera."""
    fecha = fila.get("fecha_recibido") or fila.get("fecha_publicado") or ""
    try:
        reciente = (datetime.now() - datetime.fromisoformat(fecha)).total_seconds() < 900
    except (ValueError, TypeError):
        reciente = False
    if reciente:
        logger.info(f"'{nombre}' se salteó por nombre repetido, pero la entrada previa es reciente "
                    f"(probable doble disparo); no mando aviso.")
        return
    cuando = fecha[:10] if fecha else "antes"
    cuerpo = (
        f"Subiste una {kind} en «{nombre}», pero ese nombre YA se usó para otra nota "
        f"(registrada el {cuando}). El sistema deduplica por NOMBRE, así que la salteó para no "
        f"duplicar y por eso NO se creó el borrador.\n\n"
        f"➡️ Para que salga como nota NUEVA, ponela en una carpeta con un nombre DISTINTO "
        f"(cualquiera que no hayas usado antes) y volvé a subirla.\n\n"
        f"(El nombre de la carpeta es solo interno: no aparece en la nota.)"
    )
    html = (f"<div style='font-family:Arial;max-width:600px;color:#222;font-size:16px'>"
            f"<h2 style='color:#b00020'>Nombre repetido: no se creó la nota</h2>"
            f"<p>Subiste una {kind} en «{_hesc(nombre)}», pero ese nombre <b>ya se usó</b> para otra "
            f"nota (registrada el {_hesc(cuando)}). Como el sistema deduplica por nombre, la "
            f"<b>salteó</b> para no duplicar y <b>no se creó el borrador</b>.</p>"
            f"<p>➡️ Para que salga como nota <b>nueva</b>, ponela en una carpeta con un "
            f"<b>nombre distinto</b> (cualquiera que no hayas usado) y volvé a subirla.</p>"
            f"<p style='color:#888;font-size:13px'>El nombre de la carpeta es solo interno: no "
            f"aparece en la nota.</p></div>")
    _enviar_aviso(f"Nombre repetido, no se creó la nota: {nombre}", cuerpo, html=html)


# ── ETAPA 1: preparar ─────────────────────────────────────────────────────────
def run_transcribe_video(file: str = "", uploader: str = "", dry_run: bool = False) -> None:
    modo = "SIMULACIÓN (dry-run)" if dry_run else "PROCESO REAL"
    logger.info(f"=== Desgrabar video [{modo}] — file='{file}' uploader='{uploader}' ===")

    video = _find_video(file)
    if not video:
        logger.error(f"No se encontró el video '{file}' en {_videos_folder()}.")
        return
    logger.info(f"Video: {video}")

    # #2 Corresponsal con VARIOS videos (2-3 clips cortos): SOLO el primero (por nombre) procesa el
    # grupo, concatenando todos en uno; los demás triggers saltean → varios videos = 1 sola nota.
    # `video_media` es el archivo que se le pasa a Gemini y al reel (el concatenado, o el mismo video).
    video_media = video
    grupo = []
    _ctx_g = _leer_contexto(video.parent)
    if _ctx_g and "corresponsal" in (_ctx_g.get("origen", "").lower()):
        grupo = sorted((p for p in video.parent.iterdir()
                        if p.is_file() and p.suffix.lower() in VIDEO_EXTS), key=lambda p: p.name)
    if len(grupo) > 1 and video.name != grupo[0].name:
        logger.info(f"Corresponsal multi-video: el grupo lo procesa «{grupo[0].name}»; salteo «{video.name}».")
        return

    rows = _leer_ledger()
    fila = _buscar_fila(rows, video.name)
    YA = ("borrador", "solo_reel", "publicado", "publicado_solo_reel")
    if not dry_run and fila and fila.get("estado") in YA:
        logger.info(f"El video '{video.name}' ya fue procesado (estado={fila['estado']}). Nada que hacer.")
        _avisar_nombre_repetido(video.name, fila, kind="nota de video")
        return

    WORK_DIR.mkdir(exist_ok=True)
    if len(grupo) > 1:  # soy el primero del grupo → uno todos los clips cortos en un solo video
        try:
            from video import concat_videos
            video_media = concat_videos(grupo, WORK_DIR / f"grupo_{_slug(video.stem)}.mp4")
            logger.info(f"Corresponsal: uní {len(grupo)} videos cortos en «{video_media.name}».")
        except Exception as e:  # noqa: BLE001
            logger.error(f"No pude unir el grupo de videos ({e}); proceso solo «{video.name}».")
            video_media = video
    extra_text, imgs = _recolectar_adjuntos(video)

    # Contexto del corresponsal (lo deja el bot de WhatsApp en contexto.txt). Si existe,
    # se firma el reel y se suman Lugar/Descripción al contexto que recibe Gemini.
    ctx = _leer_contexto(video.parent)
    es_corresponsal = bool(ctx and "corresponsal" in (ctx.get("origen", "").lower()))
    if ctx:
        partes = []
        if ctx.get("lugar"):
            partes.append(f"Lugar del hecho: {ctx['lugar']}")
        if ctx.get("descripcion"):
            partes.append(f"Descripción aportada por el colaborador: {ctx['descripcion']}")
        if partes:
            extra_text = (extra_text + "\n\n" + "\n".join(partes)).strip()

    # Si Gemini falla del todo (saturado/cuota tras todos los reintentos), NO tiramos abajo
    # la corrida: avisamos por mail que no se pudo esta vez y salimos limpio (exit 0). El
    # video NO se marca como procesado → se puede reintentar (re-subiéndolo o a mano).
    try:
        nota = transcribe_to_nota(video_media, extra_text=extra_text, image_paths=imgs)
    except Exception as e:
        logger.error(f"No se pudo desgrabar «{video.name}» (Gemini falló tras reintentos): {e}")
        if not dry_run:
            _enviar_aviso(
                f"No pude desgrabar el video (reintentá): {video.name}",
                f"Gemini estaba saturado y no pude desgrabar «{video.name}» esta vez.\n\n"
                f"El video NO se perdió. Para reintentarlo, volvé a subirlo a la carpeta de "
                f"«videos notas actualidad» (o avisá y lo reintento).\n\nDetalle técnico: {str(e)[:200]}",
                html=(f"<div style='font-family:Arial;max-width:600px;color:#222;font-size:16px'>"
                      f"<h2 style='color:#b00020'>No pude desgrabar el video</h2>"
                      f"<p>Gemini estaba saturado y no pude desgrabar «{_hesc(video.name)}» esta vez. "
                      f"<b>No se perdió nada.</b></p>"
                      f"<p>Para reintentarlo: volvé a subir el video a la carpeta de «videos notas "
                      f"actualidad», o avisá y lo reintento.</p>"
                      f"<p style='color:#888;font-size:13px'>Detalle: {_hesc(str(e)[:200])}</p></div>"))
        logger.info("=== Desgrabar video: fin (Gemini falló; se avisó por mail, sin marcar el video) ===")
        return

    hay = nota["hay_noticia"]
    volanta, titulo = nota["volanta"], nota["titulo"]
    texto, resumen = nota["texto"], nota["resumen"]

    # Corresponsal SIN desgrabar: si Gemini no pudo armar la nota (hay_noticia=False) pero el vecino
    # cargó una descripción en el formulario de WhatsApp, usamos ESA descripción como texto para que
    # igual salga a FB/IG/YouTube (con caption), PERO sin nota web (info no verificada, pedido del
    # usuario 2026-08-09). `corr_sin_web` fuerza hay=True (para que haya caption y salga a YouTube)
    # y evita crear el borrador de Wix (draft_id="" → en la etapa 2 la web se saltea).
    corr_sin_web = False
    if es_corresponsal and not hay and (ctx.get("descripcion") or "").strip():
        desc = ctx["descripcion"].strip()
        volanta = ""
        try:  # corregir gramática/redacción SIN cambiar la info (no reescribe/acorta/extiende)
            from utils import gemini
            corr = gemini.corregir_texto(desc, lugar=ctx.get("lugar", ""))
            titulo = corr.get("titulo") or (desc.split("\n")[0].strip()[:80] or "Envío de un corresponsal")
            texto = corr.get("texto") or desc
            resumen = corr.get("resumen") or desc[:280]
        except Exception as e:  # noqa: BLE001
            logger.warning(f"No pude corregir la descripción del vecino ({e}); la uso cruda.")
            titulo = desc.split("\n")[0].strip()[:80] or "Envío de un corresponsal"
            texto, resumen = desc, desc[:280]
        hay, corr_sin_web = True, True
        logger.info("Corresponsal sin desgrabar: uso la descripción del vecino (corregida) como texto (FB/IG/YT, sin web).")

    # Todo lo que viene DESPUÉS de la desgrabación (portada, reel, subir el reel, borrador en
    # Wix, ledger) también puede fallar por un hipo de red / GitHub Release / Wix. Si algo de
    # esto se cae, NO dejamos morir la corrida en silencio: antes el run terminaba en error, sin
    # mail y sin marcar el video → como el Apps Script ya lo marcó "visto", NO se reintentaba y
    # el video se perdía. Ahora avisamos por mail y salimos limpio (exit 0), sin marcar el video
    # → se puede reintentar re-subiéndolo. (La desgrabación de Gemini ya está protegida arriba.)
    try:
        cover = frame_at(video_media, nota["mejor_momento_seg"], WORK_DIR / "portada.jpg")
        slug = _slug(video.stem)

        # Reel para redes: VIDEO COMPLETO, sin recorte (pedido del usuario 2026-06-28).
        # Antes se recortaba a ~60s (mejores partes con noticia, o 60s sin noticia); se anuló:
        # ahora va el video entero, solo reencuadrado a vertical 9:16 para el formato reel.
        reel_path = WORK_DIR / f"reel_{slug}.mp4"
        # La firma del corresponsal ya NO se quema en el video (pedido 2026-08-05): va como TEXTO
        # al inicio de la descripción/caption en las 3 redes (se arma en run_publish_video). El
        # diario también va SIN overlay ni zócalo (2026-07-27): fondo difuminado + logo + placa.
        reel = to_vertical_reel(video_media, reel_path, overlay=False)

        if dry_run:
            logger.info(f"[dry-run] hay_noticia={hay} | tramos={len(nota.get('segmentos', []))}\n"
                        f"  VOLANTA: {volanta}\n  TÍTULO: {titulo}\n  RESUMEN: {resumen}\n  TEXTO:\n{texto}")
            logger.info(f"[dry-run] Portada: {cover}  Reel: {reel}")
            logger.info("=== Desgrabar video: fin (dry-run) ===")
            return

        reel_url = upload_reel(reel)

        draft_id = ""
        if hay and not corr_sin_web:
            # La WEB lleva el video COMPLETO (no el reel recortado): se hostea aparte y se embebe.
            web_video_url = reel_url
            try:
                full = remux_mp4(video_media, WORK_DIR / f"video_{slug}.mp4")
                web_video_url = upload_reel(full)
            except Exception as e:
                logger.warning(f"No se pudo hostear el video completo para la web ({e}); uso el reel.")
            title = f"{volanta} — {titulo}" if volanta else titulo
            body = titulo + ("\n\n" + texto if texto else "")
            info = wix.crear_borrador(title, body, cover, page=0, description=resumen, video_url=web_video_url)
            draft_id = info["draft_id"]
            estado = "borrador"
        elif corr_sin_web:
            # Borrador SOLO para poder EDITAR/BORRAR el texto desde el mail (NO se publica en la web:
            # `sin_web` lo marca). Si Wix falla, sigue sin editar/borrar (solo aprobar + previsualizar).
            try:
                body = titulo + ("\n\n" + texto if texto else "")
                info = wix.crear_borrador(titulo, body, cover, page=0, description=resumen)
                draft_id = info["draft_id"]
            except Exception as e:  # noqa: BLE001
                logger.warning(f"[wix] no pude crear el borrador de edición del corresponsal ({e}).")
            estado = "solo_reel"
        else:
            estado = "solo_reel"

        if fila is None:
            fila = {"file": video.name}
            rows.append(fila)
        fila.update({
            "uploader": uploader or fila.get("uploader", ""),
            "fecha_recibido": datetime.now().isoformat(timespec="seconds"),
            "hay_noticia": hay, "volanta": volanta, "titulo": titulo, "resumen": resumen,
            "texto": texto, "zocalo": nota.get("zocalo", ""), "draft_id": draft_id,
            "reel_url": reel_url, "estado": estado, "sin_web": corr_sin_web,
        })
        if ctx:
            fila.update({
                "origen": ctx.get("origen", ""),
                "corresponsal_nombre": ctx.get("nombre", ""),
                "corresponsal_celular": ctx.get("celular", ""),
                "corresponsal_lugar": ctx.get("lugar", ""),
                "autorizacion": ctx.get("autorizacion", ""),
            })
        _guardar_ledger(rows)
        logger.info(f"Registrado (estado={estado}, draft_id={draft_id or '—'}"
                    + (f", corresponsal={ctx.get('nombre')}" if es_corresponsal else "") + ").")
    except Exception as e:
        if dry_run:
            raise  # en pruebas manuales queremos ver el error crudo
        logger.error(f"No se pudo preparar «{video.name}» tras desgrabar (reel/subida/Wix falló): {e}")
        _enviar_aviso(
            f"No pude preparar el video (reintentá): {video.name}",
            f"Desgrabé «{video.name}» pero fallé al armar/subir el reel o crear el borrador en "
            f"Wix.\n\nEl video NO se perdió. Para reintentarlo, volvé a subirlo a la carpeta de "
            f"«videos notas actualidad» (o avisá y lo reintento).\n\nDetalle técnico: {str(e)[:200]}",
            html=(f"<div style='font-family:Arial;max-width:600px;color:#222;font-size:16px'>"
                  f"<h2 style='color:#b00020'>No pude preparar el video</h2>"
                  f"<p>Desgrabé «{_hesc(video.name)}» pero fallé al armar/subir el reel o crear el "
                  f"borrador en Wix. <b>No se perdió nada.</b></p>"
                  f"<p>Para reintentarlo: volvé a subir el video a «videos notas actualidad», o "
                  f"avisá y lo reintento.</p>"
                  f"<p style='color:#888;font-size:13px'>Detalle: {_hesc(str(e)[:200])}</p></div>"))
        logger.info("=== Desgrabar video: fin (falló la preparación; se avisó por mail, sin marcar el video) ===")
        return

    # Quién mandó el material (corresponsal de WhatsApp si lo hay; si no, el uploader de Drive).
    if es_corresponsal:
        remitente = ctx.get("nombre", "") or "corresponsal"
        if ctx.get("celular"):
            remitente += f" · {ctx['celular']}"
        if ctx.get("lugar"):
            remitente += f" · {ctx['lugar']}"
        remitente += " · Red de Corresponsales"
    else:
        remitente = uploader or "desconocido"

    if hay and not corr_sin_web:
        cuerpo = (
            f"Llegó un video para revisar: «{titulo}»\n"
            f"Enviado por: {remitente}\n\n"
            f"VOLANTA: {volanta}\nTÍTULO: {titulo}\n\nRESUMEN: {resumen}\n\n"
            f"Está cargado como BORRADOR en Wix (Blog → Borradores) con la foto de portada y el video.\n\n"
            f"➡️ Para PUBLICARLO en la web y mandar el reel a Facebook e Instagram, "
            f"mové el video «{video.name}» a la subcarpeta APROBADAS dentro de «videos notas actualidad»."
        )
        intro = (f"<h2 style='color:#e2620c'>Nota por revisar</h2>"
                 f"<p style='color:#888;font-size:13px'>{_hesc(volanta)} · enviado por {_hesc(remitente)}</p>"
                 f"<p style='font-size:19px'><b>{_hesc(titulo)}</b></p>"
                 f"<p>{_hesc(resumen)}</p>"
                 f"<p>Está como <b>borrador en Wix</b> con foto + video. Revisalo y:</p>")
        _enviar_aviso(f"Nota por revisar: {titulo}", cuerpo,
                      html=_html_aviso(intro, video.name, reel_url, draft_id, hay=True))
    elif corr_sin_web:
        cuerpo = (
            f"Llegó un video de corresponsal que NO pude desgrabar; uso la DESCRIPCIÓN que cargó el "
            f"vecino: «{titulo}»\n"
            f"Enviado por: {remitente}\n\n"
            f"TEXTO (descripción del vecino):\n{texto}\n\n"
            f"➡️ Al aprobar sale a Facebook, Instagram y YouTube con ese texto (SIN nota web). "
            f"Mové el video «{video.name}» a APROBADAS, o usá el botón «Aprobar y publicar»."
        )
        intro = (f"<h2 style='color:#e2620c'>Corresponsal sin desgrabar — sale con la descripción</h2>"
                 f"<p style='color:#888;font-size:13px'>enviado por {_hesc(remitente)}</p>"
                 f"<p style='font-size:19px'><b>{_hesc(titulo)}</b></p>"
                 f"<p style='white-space:pre-wrap'>{_hesc(texto)}</p>"
                 f"<p>No pude desgrabar el video, así que uso la descripción del vecino. Al aprobar "
                 f"sale a <b>Facebook, Instagram y YouTube</b> con ese texto (<b>sin nota web</b>):</p>")
        _enviar_aviso(f"Corresponsal por revisar (sin desgrabar): {titulo}", cuerpo,
                      html=_html_aviso(intro, video.name, reel_url, draft_id, hay=True))
    else:
        cuerpo = (
            f"Llegó un video pero NO pude desgrabarlo: «{video.name}»\n"
            f"Enviado por: {remitente}\n\n"
            f"No encontré información suficiente (ni en el audio, ni en el texto en pantalla, ni en "
            f"subtítulos o adjuntos) para armar la nota. Por eso la NOTA WEB queda SUSPENDIDA.\n\n"
            f"➡️ Si querés que igual SALGA EL REEL (recortado a 1 minuto, sin texto) a Facebook e "
            f"Instagram, mové el video «{video.name}» a la subcarpeta APROBADAS.\n"
            f"➡️ Si no, borralo. (Tip: podés re-subirlo en una subcarpeta con un .txt o fotos de "
            f"contexto para que pueda armar la nota.)"
        )
        intro = (f"<h2 style='color:#e2620c'>Video sin desgrabar</h2>"
                 f"<p>No pude armar la nota de «{_hesc(video.name)}» (no había info suficiente). "
                 f"La <b>nota web queda suspendida</b>.</p>"
                 f"<p>Si querés que igual salga <b>solo el reel</b> (1 min, sin texto):</p>")
        _enviar_aviso(f"Video sin desgrabar: {video.name}", cuerpo,
                      html=_html_aviso(intro, video.name, reel_url, "", hay=False))
    logger.info("=== Desgrabar video: fin ===")


# ── ETAPA 2: publicar (al aprobar) ────────────────────────────────────────────
def _caption(titulo: str, resumen: str) -> str:
    from utils.branding import linea_canal_yt
    site = _site()
    return (
        f"{titulo}\n\n{resumen}\n\n"
        f"📲 Seguí leyendo en {site}\n"
        f"{linea_canal_yt()}\n\n"
        f"#Chivilcoy #DiarioLaCampaña #Actualidad #Noticias"
    )


def _solo_5_hashtags(texto: str) -> str:
    """Recorta el caption para IG/FB: deja el texto hasta el 5º hashtag inclusive y borra TODO lo
    que venga después (pedido del usuario 2026-08-06). Si hay menos de 5 hashtags, lo deja igual."""
    if not texto:
        return texto
    ms = list(re.finditer(r"#[0-9A-Za-zñÑáéíóúÁÉÍÓÚ_]+", texto))
    if len(ms) < 5:
        return texto.rstrip()
    return texto[:ms[4].end()].rstrip()


# ── Aviso al corresponsal cuando su nota se publica (SOLO-GRATIS) ──────────────
_WA_CANALES = {"wix": "nuestra web", "instagram": "Instagram", "facebook": "Facebook", "youtube": "YouTube"}


def _normalizar_ar(numero: str) -> str:
    """Celular argentino: llega como 549+área+número (wa_id) y para RESPONDER va SIN el 9
    (54+área+número), igual que en el webhook (Meta lo entrega igual)."""
    n = re.sub(r"\D", "", numero or "")
    if n.startswith("549") and len(n) == 13:
        return "54" + n[3:]
    return n


def _avisar_corresponsal_publicado(celular: str, canales_ok: list, links: dict | None = None) -> None:
    """Le avisa al corresponsal por WhatsApp que su nota se publicó, con los LINKS de cada red donde
    salió (web/Instagram/Facebook/YouTube). SOLO-GRATIS: manda un mensaje LIBRE; si la ventana de
    servicio de 24 h ya cerró, Meta lo rechaza (131047) SIN costo y se saltea. NUNCA usa plantillas
    pagas. Queda DORMIDO si falta WHATSAPP_TOKEN en el .env del publicador. Nunca rompe la publicación."""
    try:
        if not celular or not canales_ok:
            return
        token = (get("WHATSAPP_TOKEN") or "").strip()
        if not token:
            logger.info("[corresponsal] aviso de publicación OMITIDO: falta WHATSAPP_TOKEN en el .env "
                        "del publicador (cargalo + sincronizá ENV_FILE para activarlo).")
            return
        phone_id = (get("WHATSAPP_PHONE_NUMBER_ID") or "1192034310668098").strip()
        to = _normalizar_ar(str(celular))
        if not to:
            return
        lista = canales_ok[0] if len(canales_ok) == 1 else ", ".join(canales_ok[:-1]) + " y " + canales_ok[-1]
        body = ("¡Hola! 🎉 Tu envío al *Programa de Corresponsales* del Diario La Campaña - Radio del "
                f"Centro ya fue *publicado* en {lista}. ¡Gracias por colaborar!")
        links = links or {}
        renglones = [f"{et} {links[k]}" for k, et in
                     (("web", "🌐 Web:"), ("instagram", "📸 Instagram:"),
                      ("facebook", "👍 Facebook:"), ("youtube", "▶️ YouTube:")) if links.get(k)]
        if renglones:
            body += "\n\n📲 Podés verlo acá:\n" + "\n".join(renglones)
        r = requests.post(
            f"https://graph.facebook.com/v21.0/{phone_id}/messages",
            headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"},
            json={"messaging_product": "whatsapp", "to": to, "type": "text", "text": {"body": body}},
            timeout=20,
        )
        if r.status_code // 100 == 2:
            logger.info(f"[corresponsal] aviso de publicación enviado a {to} (gratis, ventana 24h abierta).")
        else:
            logger.info(f"[corresponsal] aviso NO enviado (probable ventana 24h cerrada, sin costo): "
                        f"HTTP {r.status_code} {r.text[:150]}")
    except Exception as e:
        logger.warning(f"[corresponsal] no se pudo intentar el aviso de publicación: {e}")


def _retry(fn, intentos: int = 3, espera: int = 5, etiqueta: str = ""):
    """Ejecuta `fn` con reintentos automáticos (backoff lineal). Re-lanza el último
    error si agota los intentos. Se usa para YouTube y Wix (red/cuota intermitente)."""
    ultimo = None
    for i in range(max(1, intentos)):
        try:
            return fn()
        except Exception as e:  # noqa: BLE001 — queremos reintentar ante cualquier fallo de red/API
            ultimo = e
            logger.warning(f"{etiqueta or 'tarea'}: intento {i + 1}/{intentos} falló: {e}")
            if i < intentos - 1:
                time.sleep(espera * (i + 1))
    raise ultimo


# Hashtags locales SIEMPRE presentes (SEO local de Chivilcoy) + los temáticos de la nota.
_HASHTAGS_LOCALES = ["#Chivilcoy", "#NoticiasChivilcoy", "#DiarioLaCampaña", "#Actualidad"]


def _yt_enabled() -> bool:
    return (get("YT_SHORTS_ENABLED") or "1").strip().lower() not in ("0", "false", "no", "off")


def _hashtag(palabra: str) -> str:
    """Convierte 'radio del centro' → '#RadioDelCentro' (sin acentos/espacios)."""
    limpio = re.sub(r"[^0-9A-Za-zñÑáéíóúÁÉÍÓÚ ]+", "", palabra or "").strip()
    return "#" + "".join(p.capitalize() for p in limpio.split())


def _youtube_meta(volanta: str, titulo: str, resumen: str, texto: str) -> dict:
    """Arma título + descripción (formato periodístico, SEO local) + tags + hashtags del
    Short, REUTILIZANDO la lógica SEO de Gemini (`seo_youtube`). Si Gemini falla, cae a un
    armado determinístico con los datos de la nota. Nunca tira excepción."""
    site = _site()
    seo = {}
    try:
        from utils import gemini
        cuerpo_ctx = (resumen + ("\n\n" + texto if texto else "")).strip()
        seo = gemini.seo_youtube(titulo, cuerpo_ctx) or {}
    except Exception as e:  # noqa: BLE001
        logger.warning(f"[youtube] SEO con Gemini falló ({e}); uso los datos de la nota.")

    yt_titulo = (seo.get("titulo") or (f"{volanta}: {titulo}" if volanta else titulo) or titulo)[:100]

    # Hashtags: locales fijos + los temáticos que sugirió Gemini (de sus tags).
    tags = [t for t in (seo.get("tags") or []) if t]
    topicos = []
    for t in tags:
        h = _hashtag(t)
        if len(h) > 2 and h.lower() not in (x.lower() for x in _HASHTAGS_LOCALES + topicos):
            topicos.append(h)
    hashtags = _HASHTAGS_LOCALES + topicos[:6]
    linea_hashtags = " ".join(hashtags)

    # Descripción periodística: bajada (Gemini o resumen) + CTA web/suscripción + hashtags.
    bajada = (seo.get("descripcion") or resumen or titulo).strip()
    # La descripción de Gemini ya puede traer su propia línea de hashtags: la sacamos para no
    # duplicar y dejamos la nuestra (con los locales garantizados).
    bajada = re.sub(r"\n?#[\wñÑáéíóúÁÉÍÓÚ]+(?:\s+#[\wñÑáéíóúÁÉÍÓÚ]+)*\s*$", "", bajada).strip()
    descripcion = (
        f"{bajada}\n\n"
        f"📲 Seguí leyendo la nota completa en {site}\n"
        f"🔔 Suscribite al canal para más noticias de Chivilcoy y la región.\n\n"
        f"{linea_hashtags}"
    )

    # Tags de YouTube (campo Tags): los de Gemini + locales, sin '#', deduplicados.
    base_tags = ["chivilcoy", "noticias chivilcoy", "diario la campaña", "radio del centro", "actualidad"]
    final_tags, vistos = [], set()
    for t in (tags + base_tags):
        k = t.strip().lower()
        if k and k not in vistos:
            vistos.add(k)
            final_tags.append(t.strip())
    return {
        "titulo": yt_titulo,
        "descripcion": descripcion,
        "tags": final_tags[:15],
        "hashtags": linea_hashtags,
        "category_id": (get("YT_SHORTS_CATEGORY") or "25").strip(),
    }


def _meta_corresponsal(volanta: str, titulo: str, resumen: str, texto: str) -> dict:
    """Como `_youtube_meta` pero SIN reescribir el cuerpo: la descripción es el TEXTO del vecino (ya
    corregido, MISMO contenido y largo) + CTA + hashtags, y el título es el que quedó (fiel), no el
    SEO. Respeta el pedido de NO editar/acortar/extender la info del corresponsal (2026-08-09).
    Reusa `_youtube_meta` solo para los hashtags/tags (metadata, no editan la info)."""
    m = _youtube_meta(volanta, titulo, resumen, texto)
    cuerpo = (texto or resumen or titulo).strip()
    m["titulo"] = (titulo or m.get("titulo") or "")[:100]
    m["descripcion"] = f"{cuerpo}\n\n📲 Seguí leyendo en {_site()}\n\n{m['hashtags']}"
    return m


def _norm_txt(s: str) -> str:
    """Normaliza para comparar textos ignorando espacios/saltos y mayúsculas."""
    return re.sub(r"\s+", " ", (s or "")).strip().lower()


def _sincronizar_correccion(draft_id: str, volanta: str, titulo: str, texto: str,
                            resumen: str) -> tuple[str, str, str, str]:
    """Re-lee el borrador de Wix y, si el texto se CORRIGIÓ después de crearlo (botón «Corregir
    texto» o edición a mano en Wix), devuelve (volanta, titulo, texto, resumen) ACTUALIZADOS para
    que FB/IG/YouTube salgan igual que la web. Si no cambió o no se puede leer, devuelve los
    valores del ledger tal cual (sin regresión). El `resumen` (bajada) solo se regenera si cambió
    el CUERPO (no si solo se tocó el título).

    Antes la corrección vivía solo en el borrador: la web salía corregida (publica el borrador)
    pero las redes usaban el texto viejo del ledger de la etapa 1."""
    try:
        d = wix.leer_borrador_texto(draft_id)
    except Exception as e:  # noqa: BLE001
        logger.warning(f"No pude releer el borrador {draft_id} para sincronizar la corrección: {e}")
        return volanta, titulo, texto, resumen

    cuerpo = (d.get("texto") or "").strip()
    dtitle = (d.get("title") or "").strip()
    if not cuerpo:
        return volanta, titulo, texto, resumen

    esperado_cuerpo = (titulo + ("\n\n" + texto if texto else "")).strip()
    esperado_title = (f"{volanta} — {titulo}" if volanta else titulo).strip()
    cuerpo_cambio = _norm_txt(cuerpo) != _norm_txt(esperado_cuerpo)
    titulo_cambio = bool(dtitle) and _norm_txt(dtitle) != _norm_txt(esperado_title)
    if not cuerpo_cambio and not titulo_cambio:
        return volanta, titulo, texto, resumen  # sin corrección → nada que hacer

    logger.info(f"El borrador {draft_id} está corregido; sincronizo las redes con esa versión "
                f"(cuerpo_cambió={cuerpo_cambio}, título_cambió={titulo_cambio}).")

    partes = [p.strip() for p in cuerpo.split("\n\n") if p.strip()]
    # Volanta + titular: del título del post ("volanta — titular"); si no vino el título en la
    # respuesta, uso el 1er párrafo del cuerpo como titular y conservo la volanta del ledger.
    if dtitle:
        if " — " in dtitle:
            nueva_volanta, nuevo_titulo = (x.strip() for x in dtitle.split(" — ", 1))
        else:
            nueva_volanta, nuevo_titulo = "", dtitle
    else:
        nueva_volanta, nuevo_titulo = volanta, (partes[0] if partes else titulo)

    # El cuerpo del borrador arranca con el titular repetido como 1er párrafo (así se creó en la
    # etapa 1): lo saco si coincide con el titular NUEVO o el VIEJO —cubre el caso de corregir solo
    # el título—, para no duplicarlo en el caption/descripción que va a las redes. Si no coincide
    # (el usuario borró esa línea), dejo el cuerpo entero.
    if partes and _norm_txt(partes[0]) in (_norm_txt(nuevo_titulo), _norm_txt(titulo)):
        nuevo_texto = "\n\n".join(partes[1:])
    else:
        nuevo_texto = "\n\n".join(partes)

    nuevo_resumen = resumen
    if cuerpo_cambio:
        try:
            from carrusel_notas import _resumen_caption
            base = (nuevo_texto or nuevo_titulo).split("\n\n")[0]
            nuevo_resumen = _resumen_caption(base, max_chars=280) or resumen
        except Exception as e:  # noqa: BLE001
            logger.warning(f"No pude regenerar el resumen corregido ({e}); uso el del ledger.")
            nuevo_resumen = resumen

    return nueva_volanta, nuevo_titulo, nuevo_texto, nuevo_resumen


def run_publish_video(file: str = "", dry_run: bool = False) -> None:
    modo = "SIMULACIÓN (dry-run)" if dry_run else "PUBLICACIÓN REAL"
    logger.info(f"=== Publicar video aprobado [{modo}] — file='{file}' ===")

    rows = _leer_ledger()
    fila = _buscar_fila(rows, file) if file else None
    if fila is None:
        pendientes = [r for r in rows if r.get("estado") in ("borrador", "solo_reel")]
        fila = pendientes[-1] if pendientes else None
    if fila is None:
        logger.error(f"No hay nada pendiente para '{file}'. Nada que publicar.")
        return
    if fila.get("estado") in ("publicado", "publicado_solo_reel"):
        logger.info(f"El video '{fila['file']}' ya estaba publicado. Nada que hacer.")
        return

    hay = fila.get("hay_noticia", True)
    draft_id = fila.get("draft_id")
    reel_url = fila.get("reel_url")
    volanta = fila.get("volanta", "")
    titulo = fila.get("titulo", "")
    resumen = fila.get("resumen", "")
    texto = fila.get("texto", "")

    # Corrección de texto: si la nota se corrigió en el borrador de Wix (botón «Corregir texto»
    # o edición a mano) DESPUÉS de la etapa 1, esa versión vive solo en el borrador. La web sale
    # corregida (publica el borrador) pero las redes usaban el texto viejo del ledger → sincronizamos
    # volanta/título/texto/resumen para que FB/IG/YouTube salgan con la versión corregida.
    if hay and draft_id:
        volanta, titulo, texto, resumen = _sincronizar_correccion(
            draft_id, volanta, titulo, texto, resumen)

    caption = _caption(titulo, resumen) if hay else ""

    # Corresponsales: SIN firma. El caption es el TEXTO del vecino (ya corregido, NO reescrito) +
    # hashtags; en IG/FB recortado a los primeros 5 hashtags, en YouTube el texto completo. Los
    # videos del diario (no corresponsal) sí usan la bajada SEO. Nunca tira excepción.
    es_corr = "corresponsal" in (fila.get("origen", "") or "").lower()
    if hay and es_corr:
        meta = _meta_corresponsal(volanta, titulo, resumen, texto)
    elif hay and _yt_enabled():
        meta = _youtube_meta(volanta, titulo, resumen, texto)
    else:
        meta = {}
    yt_desc = meta.get("descripcion", "")
    if hay and es_corr:
        caption = yt_desc  # IG/FB parten de la descripción SEO (sin firma)
    # IG/FB: nada de texto tras los primeros 5 hashtags (en YouTube va la descripción completa).
    caption = _solo_5_hashtags(caption)

    if dry_run:
        logger.info(f"[dry-run] hay_noticia={hay}. Publicaría reel={reel_url} + draft={draft_id or '—'}\n"
                    f"Caption FB/IG:\n{caption or '(sin texto)'}\n"
                    f"YouTube Short: {'(omitido)' if not (hay and _yt_enabled()) else meta.get('titulo')}\n"
                    f"Descripción YT:\n{yt_desc or '(omitida)'}")
        return

    plats = _platforms()
    # Estado de publicación por canal (el «panel» que pide el flujo: IG/FB/YouTube/Wix).
    estado_canales = {"instagram": "omitido", "facebook": "omitido",
                      "youtube": "omitido", "wix": "omitido"}

    # Bajamos el reel UNA sola vez (lo reusan Facebook y YouTube).
    local_reel = None
    if reel_url:
        try:
            WORK_DIR.mkdir(exist_ok=True)
            local_reel = _descargar(reel_url, WORK_DIR / "reel_pub.mp4")
        except Exception as e:
            logger.error(f"No se pudo bajar el reel ({reel_url}): {e}")

    # 1) Instagram (reel remoto). Guardamos el media_id para las métricas del ranking.
    ig_media_id = fb_video_id = ""
    if "instagram" in plats and reel_url:
        try:
            res_ig = instagram.publish_reel(reel_url, caption)
            ig_media_id = (res_ig or {}).get("id", "")
            estado_canales["instagram"] = "ok"
            logger.info(f"[instagram] reel OK (id={ig_media_id})")
        except Exception as e:
            estado_canales["instagram"] = f"falló: {e}"
            logger.error(f"[instagram] reel FALLÓ: {e}")

    # 2) Facebook (archivo local).
    if "facebook" in plats and local_reel:
        try:
            res_fb = facebook.publish_video(caption, local_reel)
            fb_video_id = (res_fb or {}).get("id", "")
            estado_canales["facebook"] = "ok"
            logger.info(f"[facebook] video OK (id={fb_video_id})")
        except Exception as e:
            estado_canales["facebook"] = f"falló: {e}"
            logger.error(f"[facebook] video FALLÓ: {e}")

    # 3) YouTube Shorts (mismo archivo vertical) — solo si hay noticia (necesita SEO).
    yt_info = {}
    if hay and _yt_enabled() and local_reel:
        try:
            from platforms import youtube_api
            privacy = (get("YT_SHORTS_PRIVACY") or "public").strip()
            yt_info = _retry(
                lambda: youtube_api.upload_short(
                    local_reel, meta["titulo"], yt_desc,
                    tags=meta["tags"], category_id=meta["category_id"], privacy=privacy),
                etiqueta="[youtube] subir Short")
            estado_canales["youtube"] = "ok"
            logger.info(f"[youtube] Short OK: {yt_info.get('short_url')}")
        except Exception as e:
            estado_canales["youtube"] = f"falló: {e}"
            logger.error(f"[youtube] Short FALLÓ tras reintentos: {e}")
    elif hay and not _yt_enabled():
        logger.info("[youtube] desactivado (YT_SHORTS_ENABLED=0).")

    # 4) Nota web: embeber el YouTube (si salió) y PUBLICAR (al final del flujo). `sin_web` (corresponsal
    # sin desgrabar) tiene borrador SOLO para editar/borrar el texto → NO se publica en la web.
    post_url = ""
    if hay and draft_id and not fila.get("sin_web"):
        if yt_info.get("url"):
            try:
                _retry(lambda: wix.insertar_video_youtube(draft_id, yt_info["url"]),
                       etiqueta="[wix] embeber YouTube")
            except Exception as e:
                logger.error(f"[wix] no se pudo embeber el YouTube (la nota igual sale con el "
                             f"video nativo): {e}")
        try:
            res = _retry(lambda: wix.publicar_borrador(draft_id), etiqueta="[wix] publicar")
            post_url = res.get("url", "")
            estado_canales["wix"] = "ok"
            logger.info(f"[wix] nota publicada: {post_url}")
        except Exception as e:
            estado_canales["wix"] = f"falló: {e}"
            logger.error(f"[wix] no se pudo publicar el borrador: {e}")
    else:
        logger.info("Sin desgrabación: la nota web queda SUSPENDIDA, sale solo el reel (sin texto).")

    # Registro: estado + datos del Short (URL, ID, fecha, título; métricas a futuro).
    fila.update({
        "estado": "publicado" if hay else "publicado_solo_reel",
        "fecha_publicado": datetime.now().isoformat(timespec="seconds"),
        "post_url": post_url,
        "estado_canales": estado_canales,
        "ig_media_id": ig_media_id or fila.get("ig_media_id", ""),
        "fb_video_id": fb_video_id or fila.get("fb_video_id", ""),
        # Texto que realmente salió (ya sincronizado si se corrigió el borrador).
        "volanta": volanta, "titulo": titulo, "resumen": resumen, "texto": texto,
    })
    if yt_info:
        fila.update({
            "yt_video_id": yt_info.get("id", ""),
            "yt_url": yt_info.get("short_url") or yt_info.get("url", ""),
            "yt_watch_url": yt_info.get("watch_url", ""),
            "yt_titulo": titulo,
            "yt_privacy": yt_info.get("privacy", ""),
            "fecha_youtube": datetime.now().isoformat(timespec="seconds"),
            "yt_metrics": fila.get("yt_metrics", {}),  # se completan luego (--videos-report)
        })
    _guardar_ledger(rows)

    # Aviso de estado por canal (el «panel» de publicación).
    _avisar_estado(fila, estado_canales, post_url, yt_info)

    # Corresponsal: avisarle por WhatsApp que su nota se publicó, con los LINKS de cada red donde
    # salió ok (web/IG/FB/YouTube). SOLO-GRATIS (ver el helper).
    if es_corr:
        canales = [_WA_CANALES[k] for k in ("wix", "instagram", "facebook", "youtube")
                   if estado_canales.get(k) == "ok"]
        links = {}
        if estado_canales.get("wix") == "ok" and post_url:
            links["web"] = post_url
        if estado_canales.get("instagram") == "ok" and ig_media_id:
            links["instagram"] = instagram.permalink(ig_media_id)
        if estado_canales.get("facebook") == "ok" and fb_video_id:
            links["facebook"] = facebook.permalink(fb_video_id)
        if estado_canales.get("youtube") == "ok":
            links["youtube"] = yt_info.get("short_url") or yt_info.get("watch_url") or yt_info.get("url", "")
        _avisar_corresponsal_publicado(fila.get("corresponsal_celular", ""), canales, links)

    algun_ok = any(v == "ok" for v in estado_canales.values())
    if algun_ok:
        logger.info(f"Publicado y registrado. Estado por canal: {estado_canales}")
    else:
        logger.error("No se pudo publicar en ningún canal — revisar credenciales.")
    logger.info("=== Publicar video aprobado: fin ===")


def _placa_datos(carpeta: Path):
    """(docx, fotos, volanta, titular, resumen, texto, title, body) de una carpeta foto-nota."""
    from carrusel_notas import _resumen_caption
    docx = next((p for p in sorted(carpeta.iterdir())
                 if p.is_file() and p.suffix.lower() in (".docx", ".txt")), None)
    fotos = [p for p in sorted(carpeta.iterdir()) if p.is_file() and p.suffix.lower() in IMG_EXTS]
    if not docx or not fotos:
        return None
    volanta, titular, cuerpo = _parse_word(docx)
    if not titular:
        return None
    resumen = _resumen_caption(cuerpo[0], max_chars=280) if cuerpo else titular
    texto = "\n\n".join(cuerpo)
    title = f"{volanta} — {titular}" if volanta else titular
    body = titular + ("\n\n" + texto if texto else "")
    return docx, fotos, volanta, titular, resumen, texto, title, body


def _corresponsal_foto_etapa1(carpeta: Path, ctx: dict, uploader: str, dry_run: bool) -> None:
    """ETAPA 1 del CORRESPONSAL-FOTO: subcarpeta con contexto.txt (ORIGEN corresponsal) + foto,
    SIN video ni Word. Redacta la nota con IA desde la descripción del vecino, crea un BORRADOR en
    Wix (para poder previsualizar/corregir/borrar desde el mail) y avisa por mail para revisar. Al
    aprobar (etapa 2) se publica la nota web y se postea el reel a las redes."""
    fotos = [p for p in sorted(carpeta.iterdir())
             if p.is_file() and p.suffix.lower() in IMG_EXTS]
    if not fotos:
        logger.error(f"Corresponsal-foto '{carpeta.name}': no hay foto para procesar.")
        return
    rows = _leer_ledger()
    fila = _buscar_fila(rows, carpeta.name)
    if not dry_run and fila and fila.get("estado") in ("borrador_foto_corr", "publicado_foto_corr"):
        logger.info(f"El corresponsal-foto '{carpeta.name}' ya fue procesado (estado={fila['estado']}).")
        _avisar_nombre_repetido(carpeta.name, fila, kind="foto de corresponsal")
        return

    # IA: redactar la nota (volanta/título/texto/resumen) desde la descripción + la foto.
    from utils import gemini
    try:
        nota = gemini.nota_desde_foto(ctx.get("descripcion", ""), fotos[0], lugar=ctx.get("lugar", ""))
    except Exception as e:
        logger.error(f"Gemini falló al redactar el corresponsal-foto ({e}); uso la descripción cruda.")
        desc = (ctx.get("descripcion") or "").strip()
        nota = {"volanta": "", "titulo": (desc.split("\n")[0][:80] or "Envío de corresponsal"),
                "texto": desc, "resumen": desc[:280]}
    volanta = nota.get("volanta", ""); titular = nota.get("titulo", "")
    texto = nota.get("texto", ""); resumen = nota.get("resumen", "") or titular
    title = f"{volanta} — {titular}" if volanta else titular

    if dry_run:
        logger.info(f"[dry-run] corresponsal-foto «{title}»: nota web (borrador) + reel a redes.")
        return

    # Borrador en Wix (habilita «Corregir texto» / «Borrar» por botón + nota web al aprobar) y
    # reel de previsualización. Ambos best-effort: si Wix falla, sigue sin nota web/edición.
    draft_id = ""
    try:
        info = wix.crear_borrador_galeria(title, titular + ("\n\n" + texto if texto else ""),
                                          fotos, video_urls=[], page=0, description=resumen)
        draft_id = info["draft_id"]
    except Exception as e:  # noqa: BLE001
        logger.warning(f"[wix] no pude crear el borrador del corresponsal-foto ({e}); sigue sin nota web.")
    reel_url = _reel_preview(fotos, _slug(carpeta.name))

    if fila is None:
        fila = {"file": carpeta.name}
        rows.append(fila)
    fila.update({
        "uploader": uploader or fila.get("uploader", ""),
        "fecha_recibido": datetime.now().isoformat(timespec="seconds"),
        "hay_noticia": True, "es_placa": True, "corr_foto": True,
        "volanta": volanta, "titulo": titular, "resumen": resumen, "texto": texto,
        "draft_id": draft_id, "reel_url": reel_url, "estado": "borrador_foto_corr",
        "origen": ctx.get("origen", ""), "corresponsal_nombre": ctx.get("nombre", ""),
        "corresponsal_celular": ctx.get("celular", ""), "corresponsal_lugar": ctx.get("lugar", ""),
        "autorizacion": ctx.get("autorizacion", ""),
    })
    _guardar_ledger(rows)
    logger.info(f"Corresponsal-foto registrado como BORRADOR (draft_id={draft_id or '—'}).")

    botones = _botones_foto(carpeta.name, draft_id, reel_url)
    html = (f"<div style='font-family:Arial;max-width:600px;color:#222;font-size:16px'>"
            f"<h2 style='color:#e2620c'>Foto de corresponsal por revisar</h2>"
            f"<p style='color:#888;font-size:13px'>{_hesc(ctx.get('nombre',''))} · "
            f"{_hesc(ctx.get('lugar',''))} · foto</p>"
            f"<p style='font-size:19px'><b>{_hesc(titular)}</b></p>"
            f"<p style='white-space:pre-wrap'>{_hesc(texto or resumen)}</p>"
            f"<p>Al aprobar se publica la <b>nota web</b> y se postea el <b>reel</b> de la foto a "
            f"<b>Facebook/Instagram + YouTube Short</b> (nota redactada por IA). Con los botones "
            f"podés previsualizar el reel, corregir el texto o borrar el borrador:</p>"
            f"<div style='margin:18px 0'>{botones}</div>"
            f"<p style='color:#777;font-size:13px'>Si no ves los botones, aprobá moviendo la carpeta "
            f"«{_hesc(carpeta.name)}» a APROBADAS en Drive.</p></div>")
    _enviar_aviso(f"Foto de corresponsal por revisar: {titular}",
                  f"Foto de corresponsal para revisar: «{title}».\nPara publicarla, mové la carpeta "
                  f"«{carpeta.name}» a APROBADAS.", html=html)
    logger.info("=== Corresponsal-foto (etapa 1): fin ===")


def _corresponsal_foto_publish(fila: dict, dry_run: bool) -> None:
    """ETAPA 2 del CORRESPONSAL-FOTO (al aprobar): publica la NOTA WEB (borrador de Wix con el reel
    embebido) y postea el reel a Facebook/Instagram + YouTube Short. SIN firma. La descripción sale
    SEO con IA: completa en YouTube y recortada a 5 hashtags en IG/FB. Sincroniza la corrección de
    texto si se editó el borrador."""
    if fila.get("estado") == "publicado_foto_corr":
        logger.info(f"El corresponsal-foto '{fila['file']}' ya estaba publicado.")
        return
    carpeta = _find_folder(fila["file"])
    fotos = [p for p in sorted(carpeta.iterdir())
             if p.is_file() and p.suffix.lower() in IMG_EXTS] if carpeta else []
    if not fotos:
        logger.error(f"No encontré la foto de '{fila['file']}' para postear.")
        return
    titular = fila.get("titulo", ""); volanta = fila.get("volanta", "")
    texto = fila.get("texto", ""); resumen = fila.get("resumen", "")
    draft_id = fila.get("draft_id", "")
    # Corrección de texto: si el borrador se corrigió en Wix (botón «Corregir texto») después de la
    # etapa 1, sincronizamos para que el reel/redes Y la nota web salgan con esa versión.
    if draft_id:
        volanta, titular, texto, resumen = _sincronizar_correccion(
            draft_id, volanta, titular, texto, resumen)
    # SIN firma. Caption = TEXTO del vecino (corregido, NO reescrito) + hashtags: COMPLETO para
    # YouTube, RECORTADO a los primeros 5 hashtags para IG/FB.
    meta = _meta_corresponsal(volanta, titular, resumen, texto)
    yt_desc = meta["descripcion"]
    caption = _solo_5_hashtags(yt_desc)

    if dry_run:
        logger.info(f"[dry-run] corresponsal-foto: nota web + reel a redes (IG/FB recortado a 5 hashtags), «{titular}».")
        return

    plats = _platforms()
    estado = {"instagram": "omitido", "facebook": "omitido", "youtube": "omitido", "wix": "omitido"}
    # 1) Reel branded (30s) SIN firma quemada ni overlay del diario (la firma va en el caption).
    reel_url, reel_local = "", None
    try:
        from video import foto_a_reel
        WORK_DIR.mkdir(exist_ok=True)
        reel_local = foto_a_reel(fotos, WORK_DIR / f"corr_{_slug(fila['file'])}.mp4",
                                 overlay=False)
        reel_url = upload_reel(reel_local)
    except Exception as e:
        logger.error(f"No se pudo armar el reel del corresponsal-foto: {e}")
        return
    # 2) YouTube Short (opcional, gateado por YT_SHORTS_ENABLED). Misma descripción (firma + SEO).
    yt_info = {}
    if _yt_enabled() and reel_local:
        try:
            from platforms import youtube_api
            privacy = (get("YT_SHORTS_PRIVACY") or "public").strip()
            yt_info = _retry(lambda: youtube_api.upload_short(
                reel_local, meta["titulo"], yt_desc,
                tags=meta["tags"], category_id=meta["category_id"], privacy=privacy),
                etiqueta="[youtube] subir Short")
            estado["youtube"] = "ok"
            logger.info(f"[youtube] Short OK: {yt_info.get('short_url')}")
        except Exception as e:
            estado["youtube"] = f"falló: {e}"; logger.error(f"[youtube] Short FALLÓ: {e}")
    # 2.5) Nota web: embeber el reel (YouTube si salió) en el borrador y publicarlo.
    post_url = ""
    if draft_id:
        if yt_info.get("url"):
            try:
                _retry(lambda: wix.insertar_video_youtube(draft_id, yt_info["url"]),
                       etiqueta="[wix] embeber YouTube")
            except Exception as e:
                logger.error(f"[wix] no se pudo embeber el YouTube (la nota igual sale): {e}")
        try:
            res = _retry(lambda: wix.publicar_borrador(draft_id), etiqueta="[wix] publicar")
            post_url = (res or {}).get("url", ""); estado["wix"] = "ok"
            logger.info(f"[wix] nota publicada: {post_url}")
        except Exception as e:
            estado["wix"] = f"falló: {e}"; logger.error(f"[wix] FALLÓ: {e}")
    # 3) Reel a Instagram y Facebook con el caption.
    if "instagram" in plats and reel_url:
        try:
            res_ig = instagram.publish_reel(reel_url, caption); estado["instagram"] = "ok"
            fila["ig_media_id"] = (res_ig or {}).get("id", "") if isinstance(res_ig, dict) else ""
            logger.info("[instagram] reel OK")
        except Exception as e:
            estado["instagram"] = f"falló: {e}"; logger.error(f"[instagram] reel FALLÓ: {e}")
    if "facebook" in plats and reel_local:
        try:
            res_fb = facebook.publish_video(caption, reel_local); estado["facebook"] = "ok"
            fila["fb_video_id"] = (res_fb or {}).get("id", "") if isinstance(res_fb, dict) else ""
            logger.info("[facebook] reel OK")
        except Exception as e:
            estado["facebook"] = f"falló: {e}"; logger.error(f"[facebook] reel FALLÓ: {e}")

    rows = _leer_ledger()
    f2 = _buscar_fila(rows, fila["file"])
    if f2 is None:
        f2 = fila; rows.append(f2)
    f2.update({"estado": "publicado_foto_corr",
               "fecha_publicado": datetime.now().isoformat(timespec="seconds"),
               "estado_canales": estado, "post_url": post_url,
               "volanta": volanta, "titulo": titular, "resumen": resumen, "texto": texto,
               "ig_media_id": fila.get("ig_media_id", ""), "fb_video_id": fila.get("fb_video_id", "")})
    if yt_info:
        f2.update({"yt_video_id": yt_info.get("id", ""),
                   "yt_url": yt_info.get("short_url") or yt_info.get("url", "")})
    _guardar_ledger(rows)
    logger.info(f"Corresponsal-foto publicado: {estado}")
    _enviar_aviso(f"Corresponsal-foto publicado: {titular}",
                  f"Se publicó la foto de corresponsal «{titular}»: web={estado['wix']}, "
                  f"IG={estado['instagram']}, FB={estado['facebook']}, YT={estado['youtube']}.")
    # Aviso al corresponsal por WhatsApp con los LINKS de cada red donde salió ok (SOLO-GRATIS).
    canales = [_WA_CANALES[k] for k in ("wix", "instagram", "facebook", "youtube") if estado.get(k) == "ok"]
    links = {}
    if estado.get("wix") == "ok" and post_url:
        links["web"] = post_url
    if estado.get("instagram") == "ok" and fila.get("ig_media_id"):
        links["instagram"] = instagram.permalink(fila["ig_media_id"])
    if estado.get("facebook") == "ok" and fila.get("fb_video_id"):
        links["facebook"] = facebook.permalink(fila["fb_video_id"])
    if estado.get("youtube") == "ok":
        links["youtube"] = yt_info.get("short_url") or yt_info.get("watch_url") or yt_info.get("url", "")
    _avisar_corresponsal_publicado(fila.get("corresponsal_celular", ""), canales, links)


def run_placa(folder: str = "", uploader: str = "", dry_run: bool = False) -> None:
    """ETAPA 1 de la FOTO-NOTA: subcarpeta de «videos notas actualidad» con Word + foto(s)
    SIN video. Crea la nota como BORRADOR en Wix (la/s foto/s TAL CUAL + el texto) y avisa
    por mail para revisar. NO publica nada. Al aprobar (mover la carpeta a APROBADAS, o el
    botón) se publica la nota web y se postea la FOTO a FB/IG con TODO el texto en el caption."""
    modo = "SIMULACIÓN (dry-run)" if dry_run else "PROCESO REAL"
    logger.info(f"=== Foto-nota (etapa 1) [{modo}] — folder='{folder}' ===")
    carpeta = _find_folder(folder)
    if not carpeta:
        logger.error(f"No encontré la carpeta '{folder}' en {_videos_folder()}.")
        return
    # Corresponsal-FOTO (contexto.txt ORIGEN corresponsal + foto, SIN Word): camino aparte
    # (reel branded con firma, SOLO redes, sin nota web). Va ANTES de _placa_datos (que pide Word).
    ctx = _leer_contexto(carpeta)
    if ctx and "corresponsal" in (ctx.get("origen", "").lower()):
        _corresponsal_foto_etapa1(carpeta, ctx, uploader, dry_run)
        return
    datos = _placa_datos(carpeta)
    if not datos:
        logger.error(f"'{folder}' necesita un Word/txt con título + al menos una foto.")
        return
    _docx, fotos, volanta, titular, resumen, texto, title, body = datos

    rows = _leer_ledger()
    fila = _buscar_fila(rows, carpeta.name)
    if not dry_run and fila and fila.get("estado") in ("borrador_placa", "publicado_placa"):
        logger.info(f"La foto-nota '{carpeta.name}' ya fue procesada (estado={fila['estado']}).")
        _avisar_nombre_repetido(carpeta.name, fila, kind="foto-nota")
        return

    if dry_run:
        logger.info(f"[dry-run] foto-nota «{title}»: {len(fotos)} foto(s) tal cual + texto.\n"
                    f"  VOLANTA: {volanta}\n  TÍTULO: {titular}")
        return

    draft_id = ""
    try:
        info = wix.crear_borrador_galeria(title, body, fotos, video_urls=[], page=0, description=resumen)
        draft_id = info["draft_id"]
    except Exception as e:
        logger.error(f"[wix] no se pudo crear el borrador: {e}")
        return

    # Reel de previsualización (para el botón «Previsualizar reel» del mail). Best-effort.
    reel_url = _reel_preview(fotos, _slug(carpeta.name))

    if fila is None:
        fila = {"file": carpeta.name}
        rows.append(fila)
    fila.update({
        "uploader": uploader or fila.get("uploader", ""),
        "fecha_recibido": datetime.now().isoformat(timespec="seconds"),
        "hay_noticia": True, "es_placa": True, "volanta": volanta, "titulo": titular,
        "resumen": resumen, "texto": texto, "draft_id": draft_id, "reel_url": reel_url,
        "estado": "borrador_placa",
    })
    _guardar_ledger(rows)
    logger.info(f"Foto-nota registrada como BORRADOR (draft_id={draft_id}).")

    botones = _botones_foto(carpeta.name, draft_id, reel_url)
    html = (f"<div style='font-family:Arial;max-width:600px;color:#222;font-size:16px'>"
            f"<h2 style='color:#e2620c'>Foto-nota por revisar</h2>"
            f"<p style='color:#888;font-size:13px'>{_hesc(volanta)} · {len(fotos)} foto(s)</p>"
            f"<p style='font-size:19px'><b>{_hesc(titular)}</b></p>"
            f"<p style='white-space:pre-wrap'>{_hesc(texto or resumen)}</p>"
            f"<p>Está como <b>borrador en Wix</b> con la foto y el texto. Para PUBLICAR "
            f"(nota web con el reel + reel de la foto a Facebook/Instagram y YouTube con el texto):</p>"
            f"<div style='margin:18px 0'>{botones}</div>"
            f"<p style='color:#777;font-size:13px'>Si no ves los botones, aprobá moviendo la "
            f"carpeta «{_hesc(carpeta.name)}» a APROBADAS en Drive.</p></div>")
    _enviar_aviso(f"Foto-nota por revisar: {titular}",
                  f"Foto-nota para revisar: «{title}».\nPara publicarla, mové la carpeta "
                  f"«{carpeta.name}» a APROBADAS.", html=html)
    logger.info("=== Foto-nota (etapa 1): fin ===")


def run_placa_publish(folder: str = "", dry_run: bool = False) -> None:
    """ETAPA 2 de la FOTO-NOTA (al aprobar): arma el reel branded de la/s foto/s, lo sube al
    canal de YouTube (Diario La Campaña) con su descripción, lo embebe en la nota web y la
    publica, y postea el reel a FB/IG con TODO el texto en el caption. Mismo alcance que las
    notas de video (reel → YouTube + web con el reel + redes)."""
    modo = "SIMULACIÓN (dry-run)" if dry_run else "PUBLICACIÓN REAL"
    logger.info(f"=== Foto-nota (etapa 2: publicar) [{modo}] — folder='{folder}' ===")
    rows = _leer_ledger()
    fila = _buscar_fila(rows, folder) if folder else None
    if fila is None:
        pend = [r for r in rows if r.get("es_placa") and r.get("estado") == "borrador_placa"]
        fila = pend[-1] if pend else None
    if fila is None:
        logger.error(f"No hay foto-nota pendiente para '{folder}'.")
        return
    if fila.get("estado") == "publicado_placa":
        logger.info(f"La foto-nota '{fila['file']}' ya estaba publicada.")
        return

    # Corresponsal-FOTO: camino propio (nota web + reel a redes + aviso al vecino por WhatsApp).
    if fila.get("corr_foto"):
        _corresponsal_foto_publish(fila, dry_run)
        return

    carpeta = _find_folder(fila["file"])
    fotos = [p for p in sorted(carpeta.iterdir())
             if p.is_file() and p.suffix.lower() in IMG_EXTS] if carpeta else []
    if not fotos:
        logger.error(f"No encontré las fotos de '{fila['file']}' para postear.")
        return

    titular = fila.get("titulo", ""); volanta = fila.get("volanta", "")
    texto = fila.get("texto", ""); resumen = fila.get("resumen", "")
    draft_id = fila.get("draft_id", "")
    # Corrección de texto: si el borrador se corrigió en Wix después de la etapa 1, sincronizamos
    # para que el reel a FB/IG/YouTube salga con esa versión (la web ya sale corregida al publicar).
    if draft_id:
        volanta, titular, texto, resumen = _sincronizar_correccion(
            draft_id, volanta, titular, texto, resumen)
    title = f"{volanta} — {titular}" if volanta else titular
    site = _site()
    # TODO el texto en el caption (pedido del usuario): título + cuerpo completo + CTA.
    caption = f"{titular}\n\n{texto}\n\n📲 Seguí leyendo en {site}".strip()

    if dry_run:
        meta = _youtube_meta(volanta, titular, resumen, texto) if _yt_enabled() else {}
        logger.info(f"[dry-run] publicaría foto-nota «{title}» con {len(fotos)} foto(s): reel branded "
                    f"a FB/IG (caption completo) + YouTube Short embebido en la nota web.\n"
                    f"YouTube: {'(desactivado)' if not meta else meta.get('titulo')}\n"
                    f"Descripción YT:\n{meta.get('descripcion', '(omitida)')}")
        return

    plats = _platforms()
    # Estado por canal (mismo «panel» que las notas de video): IG/FB/YouTube/Wix.
    estado = {"instagram": "omitido", "facebook": "omitido",
              "youtube": "omitido", "wix": "omitido"}

    # 1) REEL: la/s foto/s convertida/s a un reel vertical branded (fondo naranja + logo +
    # placa de cierre), ~30 s, SIN el overlay del diario ni el texto del zócalo (pedido del
    # usuario 2026-07-27). La radio (transcriber_radio.run_placa_radio) sí lleva overlay+zócalo.
    # Se arma ANTES de publicar la web para poder subirlo a YouTube y embeberlo en la nota.
    reel_url, reel_local = "", None
    try:
        from video import foto_a_reel
        WORK_DIR.mkdir(exist_ok=True)
        reel_local = foto_a_reel(fotos, WORK_DIR / f"placa_{_slug(fila['file'])}.mp4",
                                 overlay=False)
        reel_url = upload_reel(reel_local)
    except Exception as e:
        logger.error(f"No se pudo armar el reel de la foto-nota: {e}")

    # 2) YouTube Short: el MISMO reel branded sube al canal DIARIO LA CAMPAÑA con título +
    # descripción SEO (igual que las notas de video). Reusa `_youtube_meta` (Gemini + fallback).
    yt_info = {}
    if _yt_enabled() and reel_local:
        try:
            from platforms import youtube_api
            meta = _youtube_meta(volanta, titular, resumen, texto)
            privacy = (get("YT_SHORTS_PRIVACY") or "public").strip()
            yt_info = _retry(
                lambda: youtube_api.upload_short(
                    reel_local, meta["titulo"], meta["descripcion"],
                    tags=meta["tags"], category_id=meta["category_id"], privacy=privacy),
                etiqueta="[youtube] subir Short")
            estado["youtube"] = "ok"
            logger.info(f"[youtube] Short OK: {yt_info.get('short_url')}")
        except Exception as e:
            estado["youtube"] = f"falló: {e}"
            logger.error(f"[youtube] Short FALLÓ tras reintentos: {e}")
    elif not _yt_enabled():
        logger.info("[youtube] desactivado (YT_SHORTS_ENABLED=0).")

    # 3) Nota web: embeber el reel (YouTube, si salió) DENTRO del borrador y recién PUBLICAR,
    # así la web queda «como nota con el reel» y no solo la galería de fotos.
    post_url = ""
    if draft_id:
        if yt_info.get("url"):
            try:
                _retry(lambda: wix.insertar_video_youtube(draft_id, yt_info["url"]),
                       etiqueta="[wix] embeber YouTube")
            except Exception as e:
                logger.error(f"[wix] no se pudo embeber el YouTube (la nota igual sale con las "
                             f"fotos): {e}")
        try:
            res = _retry(lambda: wix.publicar_borrador(draft_id), etiqueta="[wix] publicar")
            post_url = (res or {}).get("url", "")
            estado["wix"] = "ok"
            logger.info(f"[wix] nota publicada: {post_url}")
        except Exception as e:
            estado["wix"] = f"falló: {e}"; logger.error(f"[wix] FALLÓ: {e}")

    # 4) Reel a Instagram y Facebook con TODO el texto en el pie.
    if "instagram" in plats and reel_url:
        try:
            instagram.publish_reel(reel_url, caption); estado["instagram"] = "ok"
            logger.info("[instagram] reel OK")
        except Exception as e:
            estado["instagram"] = f"falló: {e}"; logger.error(f"[instagram] reel FALLÓ: {e}")
    if "facebook" in plats and reel_local:
        try:
            facebook.publish_video(caption, reel_local); estado["facebook"] = "ok"
            logger.info("[facebook] reel OK")
        except Exception as e:
            estado["facebook"] = f"falló: {e}"; logger.error(f"[facebook] reel FALLÓ: {e}")

    fila.update({"estado": "publicado_placa", "post_url": post_url,
                 "fecha_publicado": datetime.now().isoformat(timespec="seconds"),
                 "estado_canales": estado,
                 # Texto que realmente salió (ya sincronizado si se corrigió el borrador).
                 "volanta": volanta, "titulo": titular, "resumen": resumen, "texto": texto})
    if yt_info:
        fila.update({
            "yt_video_id": yt_info.get("id", ""),
            "yt_url": yt_info.get("short_url") or yt_info.get("url", ""),
            "yt_watch_url": yt_info.get("watch_url", ""),
            "yt_titulo": titular,
            "yt_privacy": yt_info.get("privacy", ""),
            "fecha_youtube": datetime.now().isoformat(timespec="seconds"),
        })
    _guardar_ledger(rows)

    borrar = ""
    webapp = get("APPROVE_WEBAPP_URL")
    if webapp and draft_id and estado["wix"] == "ok":
        tok = get("WEBAPP_TOKEN"); t = f"&token={quote(tok)}" if tok else ""
        borrar = (f"<div style='margin:14px 0'>"
                  f"{_boton(f'{webapp}?action=delete&post={quote(draft_id)}{t}', '🗑️ Borrar de la web', color='#b00020')}"
                  f"</div>")
    yt_url_aviso = yt_info.get("short_url") or yt_info.get("url", "")
    yt_ico = "✅" if estado["youtube"] == "ok" else ("➖" if estado["youtube"] == "omitido" else "❌")
    yt_extra = f" — <a href='{yt_url_aviso}'>{_hesc(yt_url_aviso)}</a>" if yt_url_aviso else ""
    html = (f"<div style='font-family:Arial;max-width:600px;color:#222;font-size:16px'>"
            f"<h2 style='color:#e2620c'>Foto-nota publicada</h2>"
            f"<p style='font-size:18px'><b>{_hesc(titular)}</b></p>"
            f"<ul style='line-height:1.8;list-style:none;padding:0'>"
            f"<li>{'✅' if estado['instagram'] == 'ok' else '❌'} Instagram</li>"
            f"<li>{'✅' if estado['facebook'] == 'ok' else '❌'} Facebook</li>"
            f"<li>{yt_ico} YouTube Shorts{yt_extra}</li>"
            f"<li>{'✅' if estado['wix'] == 'ok' else '❌'} Web"
            + (f" — <a href='{post_url}'>{_hesc(post_url)}</a>" if post_url else "") + "</li>"
            f"</ul>{borrar}</div>")
    _enviar_aviso(f"Foto-nota publicada: {titular}",
                  f"Se publicó «{title}» (reel de la foto a FB/IG + YouTube + nota web con el reel).\n{post_url}",
                  html=html)
    logger.info("=== Foto-nota (etapa 2): fin ===")


def _avisar_estado(fila: dict, estado: dict, post_url: str, yt_info: dict) -> None:
    """Manda un mail con el ESTADO de publicación por canal (IG/FB/YouTube/Wix)."""
    titulo = fila.get("titulo") or fila.get("file", "")
    iconos = {"ok": "✅", "omitido": "➖"}

    def _li(nombre: str, clave: str, extra: str = "") -> str:
        st = estado.get(clave, "omitido")
        ico = iconos.get(st, "❌")
        txt = "publicado" if st == "ok" else st
        return f"<li>{ico} <b>{nombre}:</b> {_hesc(txt)}{extra}</li>"

    yt_extra = ""
    if yt_info.get("short_url"):
        priv = yt_info.get("privacy", "")
        nota_priv = " <i>(privado — falta auditoría de la API)</i>" if priv and priv != "public" else ""
        yt_extra = f' — <a href="{yt_info["short_url"]}">{_hesc(yt_info["short_url"])}</a>{nota_priv}'
    wix_extra = f' — <a href="{post_url}">{_hesc(post_url)}</a>' if post_url else ""

    # Botón «Borrar de la web»: borra SOLO la nota de Wix (no FB/IG/YouTube).
    borrar = ""
    webapp = get("APPROVE_WEBAPP_URL")
    pid = fila.get("draft_id")
    if webapp and pid and estado.get("wix") == "ok":
        tok = get("WEBAPP_TOKEN")
        t = f"&token={quote(tok)}" if tok else ""
        borrar = (f"<div style='margin:14px 0'>"
                  f"{_boton(f'{webapp}?action=delete&post={quote(pid)}{t}', '🗑️ Borrar de la web', color='#b00020')}"
                  f"</div><p style='color:#999;font-size:12px'>Solo borra la nota de la web; "
                  f"el reel de FB/IG y el Short de YouTube se borran a mano.</p>")

    html = (
        f"<div style='font-family:Arial;max-width:600px;color:#222;font-size:16px'>"
        f"<h2 style='color:#e2620c'>Estado de publicación</h2>"
        f"<p style='font-size:18px'><b>{_hesc(titulo)}</b></p>"
        f"<ul style='line-height:1.8;list-style:none;padding:0'>"
        f"{_li('Instagram', 'instagram')}"
        f"{_li('Facebook', 'facebook')}"
        f"{_li('YouTube Shorts', 'youtube', yt_extra)}"
        f"{_li('Web (Wix)', 'wix', wix_extra)}"
        f"</ul>{borrar}</div>"
    )
    cuerpo = (f"Estado de publicación de «{titulo}»:\n"
              f"- Instagram: {estado.get('instagram')}\n"
              f"- Facebook: {estado.get('facebook')}\n"
              f"- YouTube: {estado.get('youtube')}"
              + (f" ({yt_info.get('short_url')})" if yt_info.get('short_url') else "") + "\n"
              f"- Web (Wix): {estado.get('wix')}" + (f" ({post_url})" if post_url else "") + "\n")
    _enviar_aviso(f"Estado de publicación: {titulo}", cuerpo, html=html)
