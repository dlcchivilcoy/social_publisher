"""Desgrabador de las notas de YouTube de Radio del Centro → texto + screenshot, por MAIL.

Toma las notas que el canal Radio del Centro subió a YouTube ese día (excluye el programa
completo «La Mañana del Centro»), las desgraba a TEXTO PERIODÍSTICO con Google Gemini
(gratis, SIN tokens de Claude, sin descargar el video: Gemini lee la URL directa), y manda
UN MAIL al diario con cada nota como Word (`.docx`: volanta + título + cuerpo) y la
miniatura (`.png`) adjuntos. (Antes dejaba los archivos en una carpeta del escritorio;
ahora van por correo a dlc.chivilcoy@gmail.com.)

Pensado para una TAREA DE WINDOWS diaria a las 14:30 (corre local). Un registro
(`.yt_desgrabaciones.json`) evita repetir.

Configuración (.env, opcional):
  GEMINI_API_KEY_YT    — clave Gemini DEDICADA a este desgrabador (su propia cuota gratis);
                         si falta, usa GEMINI_API_KEY.
  YT_DESGRABAR_EMAIL   — destinatario del mail (por defecto dlc.chivilcoy@gmail.com).
  YT_DESGRABAR_FOLDER  — carpeta de trabajo temporal (por defecto, temp del sistema).
Reusa: YT_CHANNEL_ID, STORY_EXCLUDE_TITLE, MAIL_FROM, MAIL_APP_PASSWORD, SMTP_*.
"""
import json
import mimetypes
import re
import smtplib
import ssl
import tempfile
import time
import unicodedata
from datetime import datetime
from email.message import EmailMessage
from email.utils import formataddr
from pathlib import Path

from utils.config import get
from utils.logger import get_logger

logger = get_logger("yt_desgrabador")

LEDGER = Path(__file__).parent / ".yt_desgrabaciones.json"


def _carpeta_base() -> Path:
    """Carpeta de trabajo donde se generan los .docx/.png antes de mandarlos por mail.
    Por defecto una carpeta temporal del sistema (ya NO se deja en el escritorio:
    las notas se envían por correo). Se puede forzar con YT_DESGRABAR_FOLDER."""
    raw = get("YT_DESGRABAR_FOLDER")
    if raw:
        return Path(raw)
    return Path(tempfile.gettempdir()) / "DESGRABACIONES RADIO"


def _gemini_key() -> str:
    """Clave Gemini DEDICADA al desgrabador de YouTube (su propia cuota gratis). Si no se
    cargó GEMINI_API_KEY_YT, cae a la clave general GEMINI_API_KEY."""
    return (get("GEMINI_API_KEY_YT") or "").strip() or (get("GEMINI_API_KEY") or "").strip()


def _destino_mail() -> str:
    return (get("YT_DESGRABAR_EMAIL") or "dlc.chivilcoy@gmail.com").strip()


def _enviar_por_mail(fecha: str, archivos: list[tuple[str, Path, Path | None]]) -> bool:
    """Manda UN mail con las notas desgrabadas del día como adjuntos (.docx + .png).
    `archivos` = lista de (titulo, docx_path, png_path|None). Devuelve True si se envió."""
    remitente = get("MAIL_FROM")
    password = get("MAIL_APP_PASSWORD")
    destino = _destino_mail()
    if not remitente or not password or not destino:
        logger.error("Sin credenciales de mail (MAIL_FROM/MAIL_APP_PASSWORD): no se puede enviar.")
        return False
    host = get("SMTP_HOST") or "smtp.gmail.com"
    port = int(get("SMTP_PORT") or 587)
    nombre_from = get("MAIL_FROM_NAME") or "Diario La Campaña"

    msg = EmailMessage()
    msg["From"] = formataddr((nombre_from, remitente))
    msg["To"] = destino
    msg["Subject"] = f"Desgrabaciones Radio del Centro — {fecha} ({len(archivos)} nota/s)"
    lista = "\n".join(f"• {t}" for t, _, _ in archivos)
    msg.set_content(
        f"Notas de YouTube de Radio del Centro desgrabadas el {fecha}.\n"
        f"Adjunto el Word (.docx) de cada una y la miniatura (.png).\n\n{lista}\n"
    )
    for _titulo, docx_path, png_path in archivos:
        for p in (docx_path, png_path):
            if p and p.exists():
                ctype, _ = mimetypes.guess_type(p.name)
                maintype, _, subtype = (ctype or "application/octet-stream").partition("/")
                msg.add_attachment(p.read_bytes(), maintype=maintype, subtype=subtype,
                                   filename=p.name)
    try:
        ctx = ssl.create_default_context()
        with smtplib.SMTP(host, port, timeout=120) as server:
            server.starttls(context=ctx)
            server.login(remitente, password)
            server.send_message(msg)
        logger.info(f"Mail enviado a {destino} con {len(archivos)} nota(s).")
        return True
    except Exception as e:
        logger.error(f"No se pudo enviar el mail de desgrabaciones: {e}")
        return False


def _slug(s: str, max_len: int = 80) -> str:
    """Nombre de archivo limpio: sin acentos/ñ, sin caracteres prohibidos en Windows."""
    t = unicodedata.normalize("NFD", s or "")
    t = "".join(c for c in t if unicodedata.category(c) != "Mn")
    t = t.replace("ñ", "n").replace("Ñ", "n")
    t = re.sub(r'[\\/:*?"<>|]+', " ", t)       # prohibidos en Windows
    t = re.sub(r"\s+", " ", t).strip().strip(".")
    if len(t) > max_len:
        t = t[:max_len].rsplit(" ", 1)[0]
    return t or "nota"


def _leer_ledger() -> set:
    try:
        if LEDGER.exists():
            return set(json.loads(LEDGER.read_text(encoding="utf-8")))
    except Exception:
        pass
    return set()


def _guardar_ledger(ids: set) -> None:
    LEDGER.write_text(json.dumps(sorted(ids)[-1000:], ensure_ascii=False, indent=2),
                      encoding="utf-8")


def _excluir() -> list[str]:
    """Títulos a excluir (el PROGRAMA COMPLETO). Robusto al .env: aunque
    STORY_EXCLUDE_TITLE venga con la ñ corrupta (mojibake), siempre incluye la forma sin
    acentos «manana del centro», que matchea el título normalizado del programa."""
    import youtube
    raw = get("STORY_EXCLUDE_TITLE") or ""
    items = [youtube.normalizar(x) for x in raw.split(",") if x.strip()]
    items.append("manana del centro")  # fallback que NO depende de la ñ del .env
    limpios: list[str] = []
    for x in items:
        x = re.sub(r"[^a-z0-9 ]+", "", x).strip()  # descarta mojibake/símbolos
        if x and x not in limpios:
            limpios.append(x)
    return limpios


def _png_miniatura(video_id: str, destino: Path) -> Path | None:
    """Baja la miniatura del video y la guarda como PNG en `destino`. Best-effort."""
    import youtube
    from PIL import Image
    try:
        jpg = youtube.descargar_miniatura(video_id)
        with Image.open(jpg) as im:
            im.convert("RGB").save(destino, "PNG")
        return destino
    except Exception as e:
        logger.warning(f"No se pudo guardar la miniatura PNG de {video_id}: {e}")
        return None


# LARGO de las notas del desgrabador (~2 páginas de Word). El usuario pidió (2026-07-23) que
# TODAS queden de 2 páginas: ni cortadas en menos de 1 página, ni de 3 páginas. Con el formato
# del .docx (Calibri 11, interlineado simple) 2 páginas ≈ 850–1100 palabras.
LARGO_MIN_PALABRAS = 850
LARGO_MAX_PALABRAS = 1100
LARGO_OBJETIVO = 950

INSTRUCCION_LARGO = (
    f"LARGO DE LA NOTA (REGLA FIRME): el cuerpo (campo «texto») tiene que ocupar DOS PÁGINAS "
    f"de Word — siempre 2 páginas, ni media ni tres. Apuntá a {LARGO_OBJETIVO} palabras y "
    f"quedate SIEMPRE dentro del rango {LARGO_MIN_PALABRAS}–{LARGO_MAX_PALABRAS} palabras.\n"
    "• Si el video da para MÁS: NO te pases del techo; elegí lo más importante y sintetizá el "
    "resto (no entregues notas de 3 páginas).\n"
    "• Si el video parece dar para MENOS: LLEGÁ igual a las 2 páginas desarrollando EN SERIO lo "
    "que SÍ está en el material —más contexto, antecedentes, el porqué y el para qué, citas "
    "textuales, consecuencias y próximos pasos—, SIEMPRE fiel: NO inventes, no rellenes con "
    "vueltas ni repitas la misma idea con otras palabras (no entregues notas de media página).\n"
    "• Organizá la nota en una ENTRADA que resuma lo principal, un DESARROLLO por temas y un "
    "CIERRE, con párrafos bien construidos.\n"
    "CALIDAD DE REDACCIÓN (es OBLIGATORIO, son notas que se publican con la firma del "
    "diario, tienen que quedar SERIAS y profesionales):\n"
    "• ESCRITURA CORRECTA: redactá en prosa periodística pulida, sin muletillas, sin "
    "repetir palabras pegadas ni frases. NUNCA escribas una palabra dos veces seguida "
    "(mal: «siempre siempre», «muy muy», «que que»); el reconocimiento de voz a veces "
    "duplica palabras: ELIMINÁ esas duplicaciones y reescribí la oración bien.\n"
    "• NO repitas la misma idea ni el mismo conector párrafo tras párrafo; variá el "
    "vocabulario y la construcción de las oraciones.\n"
    "• NOMBRES PROPIOS BIEN ESCRITOS: prestá MÁXIMA atención a la ortografía de nombres y "
    "apellidos de personas, lugares, clubes e instituciones. El audio puede confundir "
    "letras (B/V, C/S/Z, G/J, H muda). Si en pantalla aparece el nombre escrito (zócalo, "
    "placa, cartel), usá ESA grafía. Ante la duda, usá la forma correcta y conocida del "
    "nombre (por ejemplo, el futbolista uruguayo es «Cavani», con V). Si no podés "
    "confirmar un nombre, mejor no lo afirmes.\n"
    "• Ortografía y tildes del español rioplatense impecables; puntuación correcta."
)


_DUP_RE = re.compile(r"\b(\w+)(\s+\1\b)+", flags=re.IGNORECASE | re.UNICODE)


def _limpiar_repeticiones(texto: str) -> str:
    """Red de seguridad: colapsa palabras duplicadas pegadas que deja el reconocimiento
    de voz («siempre siempre» → «siempre», «muy muy muy» → «muy»). Conserva la primera
    aparición (con su mayúscula/acento)."""
    return _DUP_RE.sub(lambda m: m.group(1), texto or "")


def _escribir_docx(nota: dict, video: dict, path: Path) -> None:
    """Escribe la nota como Word (.docx) en formato periodístico: volanta + título +
    cuerpo en párrafos. Sin fuente ni fecha de procesado."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    if nota.get("volanta"):
        p = doc.add_paragraph()
        run = p.add_run(_limpiar_repeticiones(nota["volanta"]).upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xE2, 0x62, 0x0C)  # naranja del diario

    titulo = _limpiar_repeticiones(nota.get("titulo") or video.get("titulo", ""))
    pt = doc.add_paragraph()
    rt = pt.add_run(titulo)
    rt.bold = True
    rt.font.size = Pt(18)

    doc.add_paragraph()  # línea en blanco entre título y cuerpo
    for parrafo in _limpiar_repeticiones(nota.get("texto", "") or "").split("\n\n"):
        parrafo = parrafo.strip()
        if parrafo:
            doc.add_paragraph(parrafo)

    doc.save(str(path))


def run_yt_desgrabar(dry_run: bool = False) -> None:
    """Desgraba a texto + miniatura las notas de YouTube de hoy de Radio del Centro y las
    deja en la carpeta del escritorio. Idempotente (no repite las ya procesadas)."""
    import youtube
    from utils import gemini

    modo = "SIMULACIÓN (dry-run)" if dry_run else "PROCESO REAL"
    logger.info(f"=== Desgrabar notas de YouTube (Radio del Centro) [{modo}] ===")

    cid = get("YT_CHANNEL_ID") or "UCqiTJ2oRBLNO1ZzfrdiyjTw"
    # Por la API de datos: trae la SECCIÓN VIDEOS de hoy SIN vivos (ni vivos terminados) ni
    # shorts, y todas (no solo las 15 del RSS). Si la API falla, cae al RSS (menos preciso).
    try:
        from platforms import youtube_api
        min_seg = int(get("YT_DESGRABAR_MIN_SEG") or 60)
        videos = youtube_api.videos_seccion_de_hoy(min_seg=min_seg)
        logger.info(f"Videos de hoy (API, solo sección Videos — sin vivos ni shorts): {len(videos)}")
    except Exception as e:
        logger.warning(f"La API de YouTube falló ({e}); caigo al RSS y filtro los shorts a mano.")
        videos = youtube.videos_de_hoy(cid)
        # El RSS trae TODO (shorts incluidos): le preguntamos a YouTube uno por uno.
        try:
            from platforms import youtube_api
            videos = [v for v in videos if not youtube_api.es_short(v["id"])]
        except Exception as e2:
            logger.warning(f"Tampoco pude filtrar los shorts del RSS ({e2}).")
    if not videos:
        logger.info("No hay videos subidos hoy en el canal. Nada que desgrabar.")
        return

    excluir = _excluir()
    notas = [v for v in videos
             if not any(x and x in youtube.normalizar(v["titulo"]) for x in excluir)]
    logger.info(f"Videos de hoy: {len(videos)} | notas (sin el programa completo): {len(notas)}")

    ledger = _leer_ledger()
    pendientes = [v for v in notas if v["id"] not in ledger]
    if not pendientes:
        logger.info("Todas las notas de hoy ya estaban desgrabadas. Nada que hacer.")
        return

    fecha = datetime.now().strftime("%Y-%m-%d")
    destino = _carpeta_base() / fecha
    if not dry_run:
        destino.mkdir(parents=True, exist_ok=True)
    logger.info(f"{len(pendientes)} nota(s) por desgrabar → se enviarán por mail a {_destino_mail()}")

    key = _gemini_key()
    # Espaciar las llamadas a Gemini y REINTENTAR: la cuota gratis limita por minuto, y sin
    # esto la última nota del día solía fallar y quedar pendiente (se reintentaba recién al
    # día siguiente, donde una nota nueva la volvía a desplazar → "siempre queda una").
    entre_notas = int(get("YT_DESGRABAR_DELAY_SEG") or 8)
    reintentos = max(1, int(get("YT_DESGRABAR_REINTENTOS") or 3))
    generadas: list[tuple[str, Path, Path | None]] = []  # (titulo, docx, png) para el mail
    for i, v in enumerate(pendientes):
        if i > 0 and not dry_run:
            time.sleep(entre_notas)  # respiro entre videos para no pegarle al límite por minuto
        logger.info(f"  Desgrabando: «{v['titulo'][:60]}» ({v['url']})")
        contexto = (
            f"TÍTULO ORIGINAL DEL VIDEO EN YOUTUBE: «{v['titulo']}».\n"
            "REGLA DE ORO — NOMBRES PROPIOS Y SIGLAS (MUY IMPORTANTE: el reconocimiento de voz "
            "los equivoca MUCHO): escribí un NOMBRE PROPIO (persona, apellido, lugar, club, "
            "institución, comercio) o una SIGLA/ABREVIATURA SOLO si estás SEGURO de cómo se "
            "escribe, y estás seguro ÚNICAMENTE si figura en el TÍTULO de arriba o aparece "
            "ESCRITO en pantalla (zócalo, placa, cartel). Si NO está confirmado así, NO lo "
            "arriesgues ni adivines la grafía: reescribí la frase de forma GENÉRICA («el "
            "entrevistado», «un vecino», «una vecina», «la institución», «el club», «el "
            "funcionario», «un dirigente», «una empresa», «la organización»…) o directamente "
            "OMITÍ ese dato. Cuando el nombre SÍ está en el título o en pantalla, copialo "
            "EXACTAMENTE de ahí; si lo que escuchás en el audio no coincide, priorizá la forma "
            "del título/pantalla.\n"
            "NÚMEROS Y DATOS (cifras, montos, porcentajes, fechas, horarios, edades, resultados, "
            "cantidades): NO inventes ni redondees. Poné un número SOLO si lo escuchaste CLARO y "
            "sin dudas; ante la MÍNIMA duda, omitilo o generalizá («varios», «una parte», «cerca "
            "de»…). SIEMPRE es preferible una nota SIN ese dato que una con un nombre o un número "
            "equivocado."
        )
        # Reintenta con espera creciente ante fallos de Gemini (429/cuota/red) para que NINGUNA
        # nota quede pendiente por un límite momentáneo.
        nota = None
        for intento in range(1, reintentos + 1):
            try:
                nota = gemini.transcribe_youtube_url(v["url"], extra_text=contexto,
                                                     instrucciones=INSTRUCCION_LARGO, api_key=key)
                break
            except Exception as e:
                if intento >= reintentos:
                    logger.error(f"    Gemini falló {reintentos} veces (se reintenta en la "
                                 f"próxima corrida): {e}")
                else:
                    espera = 30 * intento
                    logger.warning(f"    Gemini falló (intento {intento}/{reintentos}): {e}. "
                                   f"Reintento en {espera}s…")
                    if not dry_run:
                        time.sleep(espera)
        if nota is None:
            continue  # NO se marca: se reintenta la próxima corrida

        if not nota.get("hay_noticia"):
            logger.info("    Sin noticia aprovechable (música/sin datos). Se saltea y se marca.")
            ledger.add(v["id"])
            if not dry_run:
                _guardar_ledger(ledger)
            continue

        # LARGO ~2 páginas: si quedó fuera de rango, un SEGUNDO PASE la ajusta (desarrolla
        # mirando de nuevo el video si quedó corta, o sintetiza si quedó larga). Así ninguna
        # nota sale de media página ni de tres páginas.
        min_pal = int(get("YT_DESGRABAR_MIN_PALABRAS") or LARGO_MIN_PALABRAS)
        max_pal = int(get("YT_DESGRABAR_MAX_PALABRAS") or LARGO_MAX_PALABRAS)
        palabras_ini = len((nota.get("texto") or "").split())
        if not dry_run and not (min_pal <= palabras_ini <= max_pal):
            estado = "corta" if palabras_ini < min_pal else "larga"
            logger.info(f"    Nota {estado} ({palabras_ini} palabras); segundo pase para dejarla "
                        f"en ~2 páginas ({min_pal}-{max_pal})…")
            nota = gemini.reescribir_a_dos_paginas(v["url"], nota, min_pal, max_pal,
                                                   LARGO_OBJETIVO, extra_text=contexto, api_key=key)

        slug = _slug(nota.get("titulo") or v["titulo"])
        docx_path = destino / f"{slug}.docx"
        png_path = destino / f"{slug}.png"
        # Evita pisar si dos notas dieran el mismo nombre
        if docx_path.exists():
            slug = f"{slug} ({v['id'][:6]})"
            docx_path = destino / f"{slug}.docx"
            png_path = destino / f"{slug}.png"

        palabras = len((nota.get("texto") or "").split())
        if dry_run:
            logger.info(f"    [dry-run] Mandaría por mail: {docx_path.name} + {png_path.name} "
                        f"(~{palabras} palabras)\n"
                        f"      VOLANTA: {nota['volanta']}\n      TÍTULO: {nota['titulo']}")
            continue

        _escribir_docx(nota, v, docx_path)
        png = _png_miniatura(v["id"], png_path)
        ledger.add(v["id"])
        _guardar_ledger(ledger)
        generadas.append((nota.get("titulo") or v["titulo"], docx_path, png))
        logger.info(f"    ✓ {docx_path.name} (~{palabras} palabras)"
                    + (f" + {png_path.name}" if png else ""))

    if dry_run:
        logger.info("=== Desgrabar notas de YouTube: fin (dry-run) ===")
        return

    if generadas:
        _enviar_por_mail(fecha, generadas)
    logger.info(f"=== Desgrabar notas de YouTube: {len(generadas)} nota(s) enviada(s) por mail ===")
